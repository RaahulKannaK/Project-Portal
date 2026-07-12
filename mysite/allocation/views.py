
import json,os,time,re
from django.shortcuts import render,get_object_or_404, redirect # type: ignore
from django.contrib import messages # type: ignore
from django.contrib.auth import logout # type: ignore
from django.http import JsonResponse # type: ignore
from .models import Student, Mentor_Login, Stu_Login,Team,Mentor,AllocationResult,Coordinator_Login
from django.conf import settings # type: ignore
from .train import allocate_mentors_ml
from django.http import JsonResponse # type: ignore
from django.views.decorators.csrf import csrf_exempt # type: ignore
from docx import Document # type: ignore
from django.http import FileResponse, Http404

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Team, ApprovedTeam, ModifyRequest,ProjectRemarks



# ---------------------
# Domain Normalization
# ---------------------
DOMAIN_MAP = {
    "AI": "Artificial Intelligence",
    "ARTIFICIAL INTELLIGENCE": "Artificial Intelligence",

    "ML": "Machine Learning",
    "MACHINE LEARNING": "Machine Learning",

    "BW": "Blockchain",
    "BLOCKCHAIN": "Blockchain",

    "CYS": "Cybersecurity",
    "CYBERSECURITY": "Cybersecurity",

    "CD": "Cloud DevOps",
    "CLOUD DEVOPS": "Cloud DevOps",

    "DS": "Data Science",
    "DATA SCIENCE": "Data Science",

    "FS": "Full Stack",
    "FULL STACK": "Full Stack",
}


def get_class_from_roll(roll):
    if roll.startswith("24UCS1"): return "CSE-A"
    elif roll.startswith("24UCS2"): return "CSE-B"
    elif roll.startswith("24UIT"): return "IT"
    elif roll.startswith("24UECE"): return "ECE"
    return "Unknown"


def normalize_domain(text: str):
    if not text:
        return ""
    cleaned = text.strip().upper()
    return DOMAIN_MAP.get(cleaned, cleaned.title())


def normalize_experience(exp: str):
    if not exp:
        return "Beginner"
    exp = exp.strip().lower()
    return "Expert" if "expert" in exp else "Beginner"


# ---------------------
# Login & Dashboards
# ---------------------
def login_view(request):
    if request.method == "POST":
        role = request.POST.get("role")
        username = request.POST.get("username")
        password = request.POST.get("password")

        if role == "student":
            try:
                login_data = Stu_Login.objects.get(username=username, password=password)
                student = Student.objects.get(student_id=username)
                request.session.update({
                    "student_id": student.student_id,
                    "student_name": student.name,
                    "student_cgpa": str(student.cgpa),
                    "student_class": get_class_from_roll(username)
                })
                return redirect("student_dashboard")
            except (Stu_Login.DoesNotExist, Student.DoesNotExist):
                messages.error(request, "Invalid student credentials or profile not found")

        elif role == "mentor":
            try:
                mentor = Mentor_Login.objects.get(username=username, password=password)
                request.session.update({
                    "mentor_id": mentor.id,
                    "username": mentor.username,
                    "mentor_name": mentor.name
                })
                return redirect("mentor_dashboard")
            except Mentor_Login.DoesNotExist:
                messages.error(request, "Invalid mentor credentials")

        elif role == "hod":
            try:
                mentor = Mentor_Login.objects.get(username=username, password=password)
                request.session.update({
                    "mentor_id": mentor.id,
                    "username": mentor.username,
                    "mentor_name": mentor.name
                })
                return redirect("hod_dashboard")
            except Mentor_Login.DoesNotExist:
                messages.error(request, "Invalid mentor credentials")
        elif role == "coordinator":
            try:
                coordinator = Coordinator_Login.objects.get(username=username, password=password)
                request.session.update({
                    "coordinator_id": coordinator.id,
                    "username": coordinator.username,
                    "mentor_name": coordinator.name
                })
                return redirect("coordinator_dashboard")
            except Coordinator_Login.DoesNotExist:
                messages.error(request, "Invalid mentor credentials")


        else:
            messages.error(request, "Invalid role selected")

    return render(request, "accounts/login.html")


from django.shortcuts import render, redirect
from django.utils import timezone
from .models import AnnouncementStatus, Stu_Login

def student_dashboard(request):
    # ===============================
    # SESSION CHECK
    # ===============================
    student_id = request.session.get("student_id")
    username = request.session.get("username")
    student_name = request.session.get("student_name")

    if not student_id:
        return redirect("login")

    password_updated = False
    password_error = None

    # ===============================
    # PASSWORD RESET HANDLER
    # ===============================
    if request.method == "POST" and request.POST.get("action") == "reset_password":
        new_password = request.POST.get("new_password")
        confirm_password = request.POST.get("confirm_password")

        if not new_password or not confirm_password:
            password_error = "Both fields are required."
        elif new_password != confirm_password:
            password_error = "Passwords do not match."
        else:
            try:
                user = Stu_Login.objects.get(username=username)
                user.password = new_password  # ⚠️ plain text as per your current setup
                user.save()
                password_updated = True
            except Stu_Login.DoesNotExist:
                password_error = "User not found."

    # ===============================
    # FETCH ANNOUNCEMENTS FOR STUDENT
    # ===============================
    announcements = AnnouncementStatus.objects.filter(
        receiver_role="student",
        receiver_id=student_id
    ).select_related("announcement").order_by("-announcement__created_at")

    # AUTO MARK AS SEEN
    announcements.filter(seen_at__isnull=True).update(
        seen_at=timezone.now()
    )

    # ===============================
    # RENDER DASHBOARD
    # ===============================
    return render(request, "student/stu_dash.html", {
        "student_name": student_name,
        "username": username,
        "student_id": student_id,
        "announcements": announcements,
        "password_updated": password_updated,
        "password_error": password_error,
    })


from django.views.decorators.http import require_POST

@require_POST
def acknowledge_announcement(request, status_id):
    status = get_object_or_404(
        AnnouncementStatus,
        id=status_id,
        receiver_role="student"
    )

    if status.acknowledged_at is None:
        status.acknowledged_at = timezone.now()
        status.save()

    return redirect("student_dashboard")

from django.shortcuts import render
from allocation.models import AllocationResult, Team, ProjectDocument, ZerothReviewRemark

def mentor_dashboard(request):
    # ===============================
    # SESSION CHECK
    # ===============================
    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")

    if not mentor_name:
        return redirect("mentor_login")

    # ===============================
    # PASSWORD RESET HANDLER
    # ===============================
    password_updated = False
    password_error = None

    if request.method == "POST" and request.POST.get("action") == "reset_password":
        new_password = request.POST.get("new_password")
        confirm_password = request.POST.get("confirm_password")

        if not new_password or not confirm_password:
            password_error = "Both fields are required."
        elif new_password != confirm_password:
            password_error = "Passwords do not match."
        else:
            try:
                user = Men_Login.objects.get(username=username)
                user.password = new_password  # ⚠️ plain text as per your current setup
                user.save()
                password_updated = True
            except Men_Login.DoesNotExist:
                password_error = "User not found."

    # ===============================
    # FETCH ALLOCATIONS & TEAM DETAILS
    # ===============================
    allocations = AllocationResult.objects.filter(mentor_name=mentor_name)
    team_details = []

    # Define review stages mapping
    review_stages = ["zeroth", "first", "second", "third"]

    for alloc in allocations:
        team = Team.objects.filter(project_title__iexact=alloc.team_name).first()
        if not team:
            continue

        members = list(zip(
            team.member_names.split(","),
            team.members.split(",")
        ))

        # ---- Documents for ALL review stages ----
        all_documents = {}
        
        for stage in review_stages:
            # Get documents where review_stage matches the stage name
            documents = ProjectDocument.objects.filter(
                team_name=team.project_title,
                review_stage=stage
            )
            doc_map = {}
            for d in documents:
                doc_map[d.doc_type] = d
            all_documents[stage] = doc_map

        # ---- Remarks for ALL review stages ----
        # Using file_type field to store review stage since no review_stage field exists
        all_remarks = {}
        
        for stage in review_stages:
            remarks = ZerothReviewRemark.objects.filter(
                team_name=team.project_title,
                mentor_name=mentor_name,
                file_type=stage  # Using existing file_type field to identify review stage
            )
            remark_map = {r.heading: r for r in remarks}
            all_remarks[stage] = remark_map

        team_details.append({
            "project_title": team.project_title,
            "domain": team.domain,
            "members": members,
            "documents": all_documents,  # All 4 stages: zeroth, first, second, third
            "remarks": all_remarks,      # All 4 stages
        })

    # ===============================
    # RENDER DASHBOARD
    # ===============================
    return render(request, "mentor/men_dash.html", {
        "mentor_name": mentor_name,
        "username": username,
        "team_details": team_details,
        "password_updated": password_updated,
        "password_error": password_error,
    })

def hod_dashboard(request):
    return render(request, "accounts/hod_dash.html")

import csv
import io
import pandas as pd
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.utils import timezone
from django.http import HttpResponse
from .models import (
    Student, Stu_Login, 
    Mentor, Mentor_Login, 
    Coordinator_Login,
    Announcement, AnnouncementStatus
)


# =========================
# HELPER: Fetch all data for template
# =========================

def get_all_data():
    """Fetch all existing data for display in coordinator dashboard"""
    
    # Students
    students_qs = Student.objects.all()
    students_data = []
    for s in students_qs:
        # Get login info if exists
        login = Stu_Login.objects.filter(username=s.student_id).first()
        students_data.append({
            'student_id': s.student_id,
            'name': s.name,
            'cgpa': s.cgpa,
            'clas': s.clas,
            'username': login.username if login else s.student_id,
            'password': login.password if login else '—'
        })
    
    # Mentors
    mentors_qs = Mentor.objects.all()
    mentors_data = []
    for m in mentors_qs:
        login = Mentor_Login.objects.filter(username=m.username).first()
        mentors_data.append({
            'username': m.username,
            'name': m.name,
            'primary_domain': m.primary_domain,
            'experience': m.experience,
            'alternative_domains': m.alternative_domains or '-',
            'password': login.password if login else '—'
        })
    
    # HODs
    hods_qs = Coordinator_Login.objects.all()
    hods_data = [{
        'username': h.username,
        'name': h.name,
        'password': h.password
    } for h in hods_qs]
    
    return {
        'students_data': students_data,
        'student_count': len(students_data),
        'mentors_data': mentors_data,
        'mentor_count': len(mentors_data),
        'hods_data': hods_data,
        'hod_count': len(hods_data),
    }


# =========================
# CSV TEMPLATE DOWNLOAD
# =========================

def download_csv_template(request, template_type):
    """Download empty CSV template with correct headers"""
    
    if template_type == 'student':
        headers = ['student_id', 'name', 'cgpa', 'clas', 'username', 'password']
        filename = 'student_template.csv'
    elif template_type == 'mentor':
        headers = ['username', 'name', 'primary_domain', 'experience', 'alternative_domains', 'password']
        filename = 'mentor_template.csv'
    elif template_type == 'hod':
        headers = ['username', 'name', 'password']
        filename = 'hod_template.csv'
    else:
        return HttpResponse("Invalid template type", status=400)
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    
    writer = csv.writer(response)
    writer.writerow(headers)
    
    return response


# =========================
# STUDENT CSV UPLOAD
# =========================

REQUIRED_STUDENT_COLUMNS = {"student_id", "name", "cgpa", "clas", "username", "password"}

def upload_student_csv(request):
    allowed = []
    not_allowed = []
    
    # Start with all existing data
    context = get_all_data()
    context['active_tab'] = 'student_csv'
    context['student_uploaded'] = 0

    if request.method == 'POST':
        file = request.FILES.get('file')

        if not file:
            context['error'] = 'No file uploaded'
            return render(request, 'coordinator/coord_dash.html', context)

        if file.name.endswith(('.xlsx', '.xls')):
            context['error'] = 'Excel files are not supported. Please upload a CSV file.'
            return render(request, 'coordinator/coord_dash.html', context)

        if not file.name.endswith('.csv'):
            context['error'] = 'Invalid file type. Only CSV files are allowed.'
            return render(request, 'coordinator/coord_dash.html', context)

        try:
            df = pd.read_csv(file)
        except Exception:
            context['error'] = 'Unable to read CSV file. Please check the format.'
            return render(request, 'coordinator/coord_dash.html', context)

        if df.empty:
            context['error'] = 'CSV file is empty.'
            return render(request, 'coordinator/coord_dash.html', context)

        # Normalize column names
        df.columns = df.columns.str.strip().str.lower()

        # Validate required columns
        missing_columns = REQUIRED_STUDENT_COLUMNS - set(df.columns)
        if missing_columns:
            context['error'] = f'Missing required columns: {", ".join(missing_columns)}. Please download the template.'
            return render(request, 'coordinator/coord_dash.html', context)

        # Validate structure
        validation_errors = []
        for idx, row in df.iterrows():
            row_num = idx + 2
            
            for col in REQUIRED_STUDENT_COLUMNS:
                if pd.isna(row.get(col)) or str(row.get(col)).strip() == '':
                    validation_errors.append(f"Row {row_num}: Missing '{col}'")
            
            try:
                cgpa = float(row.get('cgpa', 0))
                if cgpa < 0 or cgpa > 10:
                    validation_errors.append(f"Row {row_num}: CGPA must be between 0 and 10")
            except (ValueError, TypeError):
                validation_errors.append(f"Row {row_num}: Invalid CGPA value")

        if validation_errors:
            context['error'] = f"Validation failed ({len(validation_errors)} errors):<br>" + "<br>".join(validation_errors[:5])
            if len(validation_errors) > 5:
                context['error'] += f"<br>... and {len(validation_errors) - 5} more errors"
            return render(request, 'coordinator/coord_dash.html', context)

        # Process rows
        for _, row in df.iterrows():
            try:
                student_id = str(row['student_id']).strip()
                name = str(row['name']).strip()
                cgpa = float(row['cgpa'])
                clas = str(row['clas']).strip()
                username = str(row['username']).strip()
                password = str(row['password']).strip()

                # Create/Update Student
                Student.objects.update_or_create(
                    student_id=student_id,
                    defaults={
                        'name': name,
                        'cgpa': cgpa,
                        'clas': clas
                    }
                )

                # Create/Update Student Login
                Stu_Login.objects.update_or_create(
                    username=username,
                    defaults={
                        'password': password
                    }
                )

                allowed.append({
                    'student_id': student_id,
                    'name': name,
                    'cgpa': cgpa,
                    'clas': clas,
                    'username': username
                })

            except Exception as e:
                not_allowed.append({
                    'student_id': str(row.get('student_id', '')),
                    'name': str(row.get('name', '')),
                    'reason': str(e)
                })

        # Refresh ALL data after upload
        context = get_all_data()
        context.update({
            'active_tab': 'student_csv',
            'student_uploaded': len(allowed),
            'allowed': allowed,
            'not_allowed': not_allowed,
            'show_student_results': True,
            'success_msg': f'Successfully imported {len(allowed)} students!'
        })

        return render(request, 'coordinator/coord_dash.html', context)

    # GET request - just show existing data
    return render(request, 'coordinator/coord_dash.html', context)


# =========================
# MENTOR CSV UPLOAD
# =========================

REQUIRED_MENTOR_COLUMNS = {"username", "name", "primary_domain", "experience", "password"}

def upload_mentor_csv(request):
    allowed = []
    not_allowed = []
    
    # Start with all existing data
    context = get_all_data()
    context['active_tab'] = 'mentor_csv'
    context['mentor_uploaded'] = 0

    if request.method == 'POST':
        file = request.FILES.get('file')

        if not file:
            context['error'] = 'No file uploaded'
            return render(request, 'coordinator/coord_dash.html', context)

        if file.name.endswith(('.xlsx', '.xls')):
            context['error'] = 'Excel files are not supported. Please upload a CSV file.'
            return render(request, 'coordinator/coord_dash.html', context)

        if not file.name.endswith('.csv'):
            context['error'] = 'Invalid file type. Only CSV files are allowed.'
            return render(request, 'coordinator/coord_dash.html', context)

        try:
            df = pd.read_csv(file)
        except Exception:
            context['error'] = 'Unable to read CSV file. Please check the format.'
            return render(request, 'coordinator/coord_dash.html', context)

        if df.empty:
            context['error'] = 'CSV file is empty.'
            return render(request, 'coordinator/coord_dash.html', context)

        # Normalize column names
        df.columns = df.columns.str.strip().str.lower()

        # Validate required columns
        missing_columns = REQUIRED_MENTOR_COLUMNS - set(df.columns)
        if missing_columns:
            context['error'] = f'Missing required columns: {", ".join(missing_columns)}. Please download the template.'
            return render(request, 'coordinator/coord_dash.html', context)

        # Validate structure
        validation_errors = []
        for idx, row in df.iterrows():
            row_num = idx + 2
            
            for col in REQUIRED_MENTOR_COLUMNS:
                if pd.isna(row.get(col)) or str(row.get(col)).strip() == '':
                    validation_errors.append(f"Row {row_num}: Missing '{col}'")
            
            try:
                exp = int(row.get('experience', 0))
                if exp < 0:
                    validation_errors.append(f"Row {row_num}: Experience cannot be negative")
            except (ValueError, TypeError):
                validation_errors.append(f"Row {row_num}: Invalid experience value")

        if validation_errors:
            context['error'] = f"Validation failed ({len(validation_errors)} errors):<br>" + "<br>".join(validation_errors[:5])
            if len(validation_errors) > 5:
                context['error'] += f"<br>... and {len(validation_errors) - 5} more errors"
            return render(request, 'coordinator/coord_dash.html', context)

        # Process rows
        for _, row in df.iterrows():
            try:
                username = str(row['username']).strip()
                name = str(row['name']).strip()
                primary_domain = str(row['primary_domain']).strip()
                experience = int(row['experience'])
                password = str(row['password']).strip()
                alternative_domains = str(row.get('alternative_domains', '')).strip()

                # Create/Update Mentor
                Mentor.objects.update_or_create(
                    username=username,
                    defaults={
                        'name': name,
                        'primary_domain': primary_domain,
                        'experience': experience,
                        'alternative_domains': alternative_domains if alternative_domains else None
                    }
                )

                # Create/Update Mentor Login
                Mentor_Login.objects.update_or_create(
                    username=username,
                    defaults={
                        'name': name,
                        'password': password
                    }
                )

                allowed.append({
                    'username': username,
                    'name': name,
                    'primary_domain': primary_domain,
                    'experience': experience
                })

            except Exception as e:
                not_allowed.append({
                    'username': str(row.get('username', '')),
                    'name': str(row.get('name', '')),
                    'reason': str(e)
                })

        # Refresh ALL data after upload
        context = get_all_data()
        context.update({
            'active_tab': 'mentor_csv',
            'mentor_uploaded': len(allowed),
            'allowed': allowed,
            'not_allowed': not_allowed,
            'show_mentor_results': True,
            'success_msg': f'Successfully imported {len(allowed)} mentors!'
        })

        return render(request, 'coordinator/coord_dash.html', context)

    # GET request - just show existing data
    return render(request, 'coordinator/coord_dash.html', context)


# =========================
# HOD CSV UPLOAD
# =========================

REQUIRED_HOD_COLUMNS = {"username", "name", "password"}

def upload_hod_csv(request):
    allowed = []
    not_allowed = []
    
    # Start with all existing data
    context = get_all_data()
    context['active_tab'] = 'hod_csv'
    context['hod_uploaded'] = 0

    if request.method == 'POST':
        file = request.FILES.get('file')

        if not file:
            context['error'] = 'No file uploaded'
            return render(request, 'coordinator/coord_dash.html', context)

        if file.name.endswith(('.xlsx', '.xls')):
            context['error'] = 'Excel files are not supported. Please upload a CSV file.'
            return render(request, 'coordinator/coord_dash.html', context)

        if not file.name.endswith('.csv'):
            context['error'] = 'Invalid file type. Only CSV files are allowed.'
            return render(request, 'coordinator/coord_dash.html', context)

        try:
            df = pd.read_csv(file)
        except Exception:
            context['error'] = 'Unable to read CSV file. Please check the format.'
            return render(request, 'coordinator/coord_dash.html', context)

        if df.empty:
            context['error'] = 'CSV file is empty.'
            return render(request, 'coordinator/coord_dash.html', context)

        # Normalize column names
        df.columns = df.columns.str.strip().str.lower()

        # Validate required columns
        missing_columns = REQUIRED_HOD_COLUMNS - set(df.columns)
        if missing_columns:
            context['error'] = f'Missing required columns: {", ".join(missing_columns)}. Please download the template.'
            return render(request, 'coordinator/coord_dash.html', context)

        # Validate structure
        validation_errors = []
        for idx, row in df.iterrows():
            row_num = idx + 2
            
            for col in REQUIRED_HOD_COLUMNS:
                if pd.isna(row.get(col)) or str(row.get(col)).strip() == '':
                    validation_errors.append(f"Row {row_num}: Missing '{col}'")

        if validation_errors:
            context['error'] = f"Validation failed ({len(validation_errors)} errors):<br>" + "<br>".join(validation_errors[:5])
            if len(validation_errors) > 5:
                context['error'] += f"<br>... and {len(validation_errors) - 5} more errors"
            return render(request, 'coordinator/coord_dash.html', context)

        # Process rows
        for _, row in df.iterrows():
            try:
                username = str(row['username']).strip()
                name = str(row['name']).strip()
                password = str(row['password']).strip()

                # Create/Update HOD
                Coordinator_Login.objects.update_or_create(
                    username=username,
                    defaults={
                        'name': name,
                        'password': password
                    }
                )

                allowed.append({
                    'username': username,
                    'name': name
                })

            except Exception as e:
                not_allowed.append({
                    'username': str(row.get('username', '')),
                    'name': str(row.get('name', '')),
                    'reason': str(e)
                })

        # Refresh ALL data after upload
        context = get_all_data()
        context.update({
            'active_tab': 'hod_csv',
            'hod_uploaded': len(allowed),
            'allowed': allowed,
            'not_allowed': not_allowed,
            'show_hod_results': True,
            'success_msg': f'Successfully imported {len(allowed)} HODs!'
        })

        return render(request, 'coordinator/coord_dash.html', context)

    # GET request - just show existing data
    return render(request, 'coordinator/coord_dash.html', context)


# =========================
# COORDINATOR DASHBOARD
# =========================

def coordinator_dashboard(request):
    coordinator_id = request.session.get("coordinator_id")
    if not coordinator_id:
        return redirect("login")

    coordinator = get_object_or_404(Coordinator_Login, id=coordinator_id)

    # Password reset handler
    password_updated = False
    password_error = None

    if request.method == "POST" and request.POST.get("action") == "reset_password":
        new_password = request.POST.get("new_password")
        confirm_password = request.POST.get("confirm_password")

        if not new_password or not confirm_password:
            password_error = "Both fields are required."
        elif new_password != confirm_password:
            password_error = "Passwords do not match."
        else:
            try:
                user = Coordinator_Login.objects.get(username=coordinator.username)
                user.password = new_password
                user.save()
                password_updated = True
            except Coordinator_Login.DoesNotExist:
                password_error = "User not found."

    # Create announcement
    if request.method == "POST" and 'title' in request.POST:
        title = request.POST.get("title")
        ann_type = request.POST.get("ann_type")
        target = request.POST.get("target")

        deadline_date = request.POST.get("deadline_date")
        deadline_time = request.POST.get("deadline_time")

        schedule_date = request.POST.get("schedule_date")
        schedule_time = request.POST.get("schedule_time")
        venue = request.POST.get("venue")

        message = request.POST.get("message")

        if not title or not ann_type or not target:
            messages.error(request, "Please fill all required fields")
            return redirect("coordinator_dashboard")

        announcement = None

        if ann_type == "deadline":
            if not deadline_date or not deadline_time:
                messages.error(request, "Deadline date and time required")
                return redirect("coordinator_dashboard")

            announcement = Announcement.objects.create(
                title=title,
                ann_type=ann_type,
                target_role=target,
                deadline_date=deadline_date,
                deadline_time=deadline_time,
                created_by_username=coordinator.username,
                created_by_name=coordinator.name
            )

        elif ann_type == "schedule":
            if not schedule_date or not schedule_time or not venue:
                messages.error(request, "Schedule date, time and venue required")
                return redirect("coordinator_dashboard")

            announcement = Announcement.objects.create(
                title=title,
                ann_type=ann_type,
                target_role=target,
                schedule_date=schedule_date,
                schedule_time=schedule_time,
                venue=venue,
                created_by_username=coordinator.username,
                created_by_name=coordinator.name
            )

        elif ann_type == "instruction":
            if not message:
                messages.error(request, "Instruction message required")
                return redirect("coordinator_dashboard")

            announcement = Announcement.objects.create(
                title=title,
                ann_type=ann_type,
                target_role=target,
                message=message,
                created_by_username=coordinator.username,
                created_by_name=coordinator.name
            )

        else:
            messages.error(request, "Invalid announcement type")
            return redirect("coordinator_dashboard")

        status_objects = []

        if target == "student":
            students = Student.objects.all()
            for s in students:
                status_objects.append(
                    AnnouncementStatus(
                        announcement=announcement,
                        receiver_role="student",
                        receiver_id=s.student_id,
                        receiver_name=s.name
                    )
                )

        elif target == "mentor":
            mentors = Mentor_Login.objects.all()
            for m in mentors:
                status_objects.append(
                    AnnouncementStatus(
                        announcement=announcement,
                        receiver_role="mentor",
                        receiver_id=m.username,
                        receiver_name=m.name
                    )
                )

        else:
            students = Student.objects.all()
            mentors = Mentor_Login.objects.all()

            for s in students:
                status_objects.append(
                    AnnouncementStatus(
                        announcement=announcement,
                        receiver_role="student",
                        receiver_id=s.student_id,
                        receiver_name=s.name
                    )
                )

            for m in mentors:
                status_objects.append(
                    AnnouncementStatus(
                        announcement=announcement,
                        receiver_role="mentor",
                        receiver_id=m.username,
                        receiver_name=m.name
                    )
                )

        AnnouncementStatus.objects.bulk_create(status_objects)

        messages.success(request, "Announcement circulated successfully!")
        return redirect("coordinator_dashboard")

    # Load dashboard data - ALL DATA including lists
    context = get_all_data()
    context.update({
        "coordinator": coordinator,
        "announcements": Announcement.objects.filter(
            created_by_username=coordinator.username
        ).order_by("-created_at"),
        "password_updated": password_updated,
        "password_error": password_error,
    })

    return render(request, "coordinator/coord_dash.html", context)

def logout_view(request):
    logout(request)
    return redirect("login")


@csrf_exempt
def team_partitions(total_students, team_sizes=(3, 4)):
    """Generate all possible team formations with teams of 3 and 4 members."""
    results = []

    def backtrack(remaining, current):
        if remaining == 0:
            results.append(list(current))
            return
        for size in team_sizes:
            if remaining - size >= 0:
                current.append(size)
                backtrack(remaining - size, current)
                current.pop()

    backtrack(total_students, [])

    # Normalize by counts (ignore order)
    unique = []
    final = []
    for combo in results:
        counts = (combo.count(3), combo.count(4))  # (teams of 3, teams of 4)
        if counts not in unique:
            unique.append(counts)
            final.append(counts)

    return final


def find_best_possibilities(possibilities):
    """
    Best definition (neutral & fair):
    - fewer total teams is better
    - 3 and 4 member teams are both allowed
    """
    # total teams = teams_of_3 + teams_of_4
    min_teams = min(a + b for a, b in possibilities)

    best = []
    for a, b in possibilities:
        if a + b == min_teams:
            best.append((a, b))

    return best


def calculate_team_possibilities(student_class):
    """
    Calculate team possibilities for available students in a class.
    Returns dict with total available, possibilities, and best options.
    """
    # Get all students in class
    all_class_students = Student.objects.filter(clas=student_class)
    total_in_class = all_class_students.count()
    
    # Get used students (already in teams)
    used_rolls = []
    for t in Team.objects.filter(student_class=student_class):
        if t.members:
            used_rolls.extend(t.members.split(","))
    
    # Calculate available students
    available_students = total_in_class - len(set(used_rolls))
    
    # Get possibilities if enough students
    possibilities = []
    best_possibilities = []
    
    if available_students >= 3:
        possibilities = team_partitions(available_students)
        best_possibilities = find_best_possibilities(possibilities)
    
    return {
        "total_in_class": total_in_class,
        "used_students": len(set(used_rolls)),
        "available_students": available_students,
        "possibilities": possibilities,
        "best_possibilities": best_possibilities,
        "can_form_teams": available_students >= 3
    }


import json
from django.http import JsonResponse
from django.shortcuts import render

def create_team(request):
    student_class = request.session.get("student_class")
    student_id = request.session.get("student_id")

    # Calculate team possibilities for the class
    team_data = calculate_team_possibilities(student_class)
    
    # Collect all members already used in any team (global check)
    used_rolls = []
    for t in Team.objects.all():
        if t.members:
            used_rolls.extend(t.members.split(","))

    # Find existing team for this student
    existing_team = None
    for t in Team.objects.filter(student_class=student_class):
        if t.members and student_id in t.members.split(","):
            existing_team = t
            break
    
    already_created = existing_team is not None
    
    # If updating, exclude current team members from used_rolls so they can be reselected
    if existing_team and existing_team.members:
        current_members = existing_team.members.split(",")
        used_rolls = [r for r in used_rolls if r not in current_members]
    
    classmates = Student.objects.filter(clas=student_class).exclude(student_id__in=used_rolls)

    # -------------------------
    # Handle form submission (POST)
    # -------------------------
    if request.method == "POST":
        try:
            body = request.body.decode("utf-8")
            data = json.loads(body) if body else {}

            print("Received data:", data)

            project_title = str(data.get("project_title", "")).strip()
            domain_raw = data.get("domain", "")
            domain = domain_raw.upper() if domain_raw else ""
            members = data.get("members")  # Can be None, list, or null
            is_update = data.get("is_update", False)

            # Validation checks
            if not project_title:
                return JsonResponse({"status": "error", "message": "Project title is required."})

            if not domain:
                return JsonResponse({"status": "error", "message": "Domain is required."})

            # Determine if this is a member update or just title/domain update
            is_member_update = members is not None and len(members) > 0

            # For updates where members aren't being changed, preserve existing members
            if is_update and existing_team and not is_member_update:
                # Use existing members from database
                all_members = set(existing_team.members.split(",")) if existing_team.members else set()
                
                # Ensure leader is included (should already be there)
                if student_id not in all_members:
                    return JsonResponse({
                        "status": "error",
                        "message": "You must be part of the team!"
                    })
            else:
                # New team or explicit member update
                all_members = set(members) if members else set()
                all_members.add(student_id)

                # Check if logged-in student is part of the submitted members
                if student_id not in all_members:
                    return JsonResponse({
                        "status": "error",
                        "message": "You must be part of the team!"
                    })

                # For new teams: Check if any members are already used
                if not is_update:
                    already_used = [m for m in all_members if m in used_rolls]
                    if already_used:
                        return JsonResponse({
                            "status": "error",
                            "message": f"Student(s) {', '.join(already_used)} already in a team."
                        })

                # Validate team size (3 or 4 members) only when members are provided
                if len(all_members) not in [3, 4]:
                    return JsonResponse({
                        "status": "error",
                        "message": f"Team must have 3 or 4 members. You selected {len(all_members)}."
                    })

            # Check for duplicate project titles (exclude self if updating)
            title_qs = Team.objects.filter(project_title__iexact=project_title)
            if is_update and existing_team:
                title_qs = title_qs.exclude(id=existing_team.id)
            
            if title_qs.exists():
                return JsonResponse({
                    "status": "error",
                    "message": "Project title already exists. Please choose another."
                })

            # Get names of selected members
            member_objs = Student.objects.filter(student_id__in=all_members)
            member_names = [s.name for s in member_objs]

            if is_update and existing_team:
                # UPDATE existing team
                existing_team.project_title = project_title
                existing_team.domain = domain
                
                # Only update members if explicitly provided
                if is_member_update:
                    existing_team.members = ",".join(sorted(all_members))
                    existing_team.member_names = ",".join(member_names)
                    # Keep original leader - don't change leader on member updates
                
                # CLEAR all update flags — back to normal pending state
                existing_team.needs_update_problem = False
                existing_team.needs_update_domain = False
                existing_team.needs_update_members = False
                existing_team.modification_reason = ''
                existing_team.status = 'pending'  # Back to pending for coordinator review
                existing_team.save()
                
                return JsonResponse({
                    "status": "success", 
                    "project_title": project_title, 
                    "message": "Team updated successfully"
                })
            else:
                # CREATE new team - leader is the creator (logged-in student)
                Team.objects.create(
                    project_title=project_title,
                    student_class=student_class,
                    domain=domain,
                    members=",".join(sorted(all_members)),
                    member_names=",".join(member_names),
                    leader_id=student_id,  # Store who created the team
                    status='pending'
                )

                return JsonResponse({"status": "success", "project_title": project_title})

        except json.JSONDecodeError:
            return JsonResponse({"status": "error", "message": "Invalid JSON data received."})
        except Exception as e:
            print("Exception in create_team:", e)
            return JsonResponse({"status": "error", "message": str(e)})

    # -------------------------
    # Prepare page data for render
    # -------------------------
    members_list = []
    leader_id = None
    
    if existing_team and existing_team.members:
        ids = existing_team.members.split(",")
        # Get leader_id from model, or default to first member for legacy teams
        leader_id = getattr(existing_team, 'leader_id', None) or ids[0]
        members_list = list(
            Student.objects.filter(student_id__in=ids).values_list("student_id", "name")
        )

    return render(request, "student/create_team.html", {
        "classmates": classmates,
        "student_class": student_class,
        "already_created": already_created,
        "existing_team": existing_team,
        "members_list": members_list,
        "leader_id": leader_id,  # Pass leader to template
        "team_data": team_data,
        "possibilities": team_data["possibilities"],
        "best_possibilities": team_data["best_possibilities"],
        "available_students": team_data["available_students"],
        "can_form_teams": team_data["can_form_teams"],
    })

def view_mentor(request):
    return render(request, "accounts/view_mentor.html")

def add_men(request):
    if "mentor_id" not in request.session:
        return JsonResponse({"status": "error", "message": "Not logged in"}, status=401)

    mentor_created = False
    primary_domain = experience = alternative_domains_list = ""
    username = mentor_name = None

    username = request.session.get("username")
    mentor_name = request.session.get("mentor_name")

    try:
        mentor = Mentor.objects.get(username=username)
        mentor_created = True
        primary_domain = mentor.primary_domain
        experience = mentor.experience
        # Split alternative domains into list for template
        alternative_domains_list = mentor.alternative_domains.split(",") if mentor.alternative_domains else []
    except Mentor.DoesNotExist:
        pass

    if request.method == "POST":
        try:
            data = json.loads(request.body)
            primary_domain = data.get("primary_domain")
            experience = data.get("experience")
            alt_domains = data.get("alt_domains", [])
            alternative_domains = ",".join(alt_domains)

            Mentor.objects.update_or_create(
                username=username,
                defaults={
                    "name": mentor_name,
                    "primary_domain": primary_domain,
                    "experience": experience,
                    "alternative_domains": alternative_domains
                }
            )

            return JsonResponse({"status": "success"})

        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=500)

    context = {
        "mentor_created": mentor_created,
        "primary_domain": primary_domain,
        "experience": experience,
        "alternative_domains": alternative_domains_list,
        "username": username,
        "mentor_name": mentor_name
    }
    return render(request, "mentor/add_men.html", context)

import json
from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

from .models import Team, Mentor, AllocationResult


def allocate_view(request):
    # ---------------------
    # Fetch teams from DB
    # ---------------------
    teams_qs = Team.objects.all()
    teams = [{"id": t.id, "name": t.project_title, "domain": t.domain} for t in teams_qs]
    print("✅ Teams fetched from DB:", teams)

    if not teams:
        return render(request, "coordinator/men_team_result.html", {
            "allocations": [],
            "teams_data": [],
            "mentors_data": [],
            "avg_similarity": 0,
            "error": "No teams available."
        })

    # ---------------------
    # Fetch mentors from DB
    # ---------------------
    mentors_qs = Mentor.objects.all()
    mentors = []
    mentors_data = []  # For JS autocomplete in template
    
    for m in mentors_qs:
        alt_list = [d.strip() for d in m.alternative_domains.split(",")] if m.alternative_domains else []
        
        mentors.append({
            "id": m.id,
            "domain": m.primary_domain,
            "name": m.name,
            "experience": "Expert" if m.experience >= 4 else "Beginner",
            "alt_domains": alt_list
        })
        
        mentors_data.append({
            "id": m.id,
            "name": m.name,
            "domain": m.primary_domain,
            "alt_domains": json.dumps(alt_list)
        })
    
    print("✅ Mentors fetched from DB:", mentors)

    if not mentors:
        return render(request, "coordinator/men_team_result.html", {
            "allocations": [],
            "teams_data": [],
            "mentors_data": [],
            "avg_similarity": 0,
            "error": "No mentors available."
        })

    # ---------------------
    # ML Allocation
    # ---------------------
    allocations_df = allocate_mentors_ml(teams, mentors)
    print("✅ Allocations DataFrame:\n", allocations_df)

    # ---------------------
    # Rename columns for template & DB
    # ---------------------
    allocations_df.rename(columns={
        "Team": "team_name",
        "Team Domain": "team_domain",
        "Mentor": "mentor_name",
        "Mentor Domain": "mentor_domain",
        "Mentor Alt Domains": "alt_domains",
        "Experience": "experience",
        "Similarity Score": "similarity_score",
        "Reason": "reason"
    }, inplace=True)

    # Convert alt_domains to comma-separated strings
    allocations_df["alt_domains"] = allocations_df["alt_domains"].fillna("").apply(
        lambda x: x if isinstance(x, str) else ", ".join(x)
    )

    # Calculate average similarity
    avg_similarity = round(allocations_df["similarity_score"].mean(), 1) if not allocations_df.empty else 0

    # ---------------------
    # Save allocations to DB
    # ---------------------
    for _, row in allocations_df.iterrows():
        AllocationResult.objects.update_or_create(
            team_name=row["team_name"],
            defaults={
                "team_domain": row["team_domain"],
                "mentor_name": row["mentor_name"],
                "mentor_domain": row["mentor_domain"],
                "alt_domains": row["alt_domains"],
                "experience": row["experience"],
                "similarity_score": row["similarity_score"],
                "reason": row["reason"]
            }
        )

    # ---------------------
    # Prepare template data
    # ---------------------
    allocations = allocations_df.to_dict(orient="records")
    print("✅ Allocations list for template:", allocations)

    # Teams data for JS autocomplete
    teams_data = [{
        "id": t.id,
        "name": t.project_title,
        "domain": t.domain
    } for t in teams_qs]

    return render(request, "coordinator/men_team_result.html", {
        "allocations": allocations,
        "teams_data": teams_data,
        "mentors_data": mentors_data,
        "avg_similarity": avg_similarity,
    })


@csrf_exempt
def save_allocations(request):
    """
    Save manually edited allocations from Table 2
    """
    if request.method != "POST":
        return JsonResponse({"status": "fail", "message": "POST required"}, status=400)

    try:
        data = json.loads(request.body)
        allocations = data.get("allocations", [])
        
        print(f"💾 Saving {len(allocations)} allocations...")

        # Option 1: Replace all existing allocations
        # AllocationResult.objects.all().delete()
        
        # Option 2: Update or create each
        for idx, alloc in enumerate(allocations):
            team_name = alloc.get("teamName", "").strip()
            if not team_name:
                continue
                
            AllocationResult.objects.update_or_create(
                team_name=team_name,
                defaults={
                    "team_domain": alloc.get("teamDomain", ""),
                    "mentor_name": alloc.get("mentorName", ""),
                    "mentor_domain": alloc.get("mentorDomain", ""),
                    "experience": "Unknown",  # Or derive from mentor lookup
                    "similarity_score": 0,     # Manual entry has no ML score
                    "reason": "Manually assigned",
                    "alt_domains": ""
                }
            )

        return JsonResponse({
            "status": "success",
            "saved": len(allocations),
            "message": f"{len(allocations)} allocations saved successfully"
        })

    except Exception as e:
        print("❌ Save error:", e)
        import traceback
        traceback.print_exc()
        return JsonResponse({
            "status": "fail",
            "message": str(e)
        }, status=500)

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.db.models import Avg
from .models import Team, Mentor, AllocationResult

@login_required
def student_result_view(request):
    """
    View for students to see their assigned mentor.
    Fetches the team associated with the logged-in student and shows mentor details.
    """
    try:
        # Get the team where current user is a member
        # Adjust this query based on your Team model's relationship to User
        team = Team.objects.filter(members=request.user).first() or \
               Team.objects.filter(leader=request.user).first() or \
               Team.objects.filter(user=request.user).first()
        
        if not team:
            return render(request, "student/student_result.html", {
                "allocation": None,
                "error": "You are not associated with any team."
            })
        
        # Get allocation for this team
        allocation = AllocationResult.objects.filter(team_name=team.project_title).first()
        
        if not allocation:
            return render(request, "student/student_result.html", {
                "allocation": None,
                "error": "Mentor allocation not found for your team."
            })
        
        context = {
            "allocation": allocation,
            "team": team
        }
        
        return render(request, "student/student_result.html", context)
        
    except Exception as e:
        return render(request, "student/student_result.html", {
            "allocation": None,
            "error": str(e)
        })


@login_required
def mentor_result_view(request):
    """
    View for mentors to see their assigned teams.
    Fetches all allocations where the logged-in mentor is assigned.
    """
    try:
        # Get mentor profile for current user
        # Adjust this query based on your Mentor model's relationship to User
        mentor = Mentor.objects.filter(user=request.user).first() or \
                 Mentor.objects.filter(email=request.user.email).first() or \
                 Mentor.objects.filter(name=request.user.get_full_name()).first()
        
        if not mentor:
            return render(request, "mentor/mentor_result.html", {
                "allocations": [],
                "error": "Mentor profile not found."
            })
        
        # Get all allocations for this mentor
        allocations = AllocationResult.objects.filter(mentor_name=mentor.name)
        
        # Calculate statistics
        unique_domains = allocations.values_list('team_domain', flat=True).distinct()
        avg_similarity = allocations.aggregate(Avg('similarity_score'))['similarity_score__avg'] or 0
        
        context = {
            "allocations": allocations,
            "unique_domains": list(unique_domains),
            "avg_similarity": round(avg_similarity, 1),
            "active_count": allocations.count(),
            "mentor": mentor
        }
        
        return render(request, "mentor/mentor_result.html", context)
        
    except Exception as e:
        return render(request, "mentor/mentor_result.html", {
            "allocations": [],
            "error": str(e)
        })

# views.py

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.db.models import Avg
from django.conf import settings
from .models import Team, Mentor, AllocationResult, Stu_Login, Mentor_Login

def get_logged_in_student(request):
    """Helper to get student info from session"""
    student_id = request.session.get('student_id')
    student_name = request.session.get('student_name')
    student_class = request.session.get('student_class')
    return {
        'id': student_id,
        'name': student_name,
        'class': student_class
    }

def get_logged_in_mentor(request):
    """Helper to get mentor info from session - uses name only"""
    # Your session has: ['mentor_id', 'mentor_name'] (no username)
    mentor_id = request.session.get('mentor_id')
    mentor_name = request.session.get('mentor_name')
    
    print("Session keys:", list(request.session.keys()))
    print("Mentor ID from session:", mentor_id)
    print("Mentor name from session:", mentor_name)
    
    return {
        'id': mentor_id,
        'name': mentor_name
    }

def student_result_view(request):
    """
    View for students to see their assigned mentor.
    Uses session data from your custom login system.
    """
    # Check if student is logged in (custom auth)
    student_info = get_logged_in_student(request)
    
    if not student_info['id']:
        # Try Django's auth as fallback
        if not request.user.is_authenticated:
            from django.shortcuts import redirect
            return redirect('/login/')
        # If using Django auth, get student by username
        student_name = request.user.get_full_name() or request.user.username
    else:
        student_name = student_info['name']
    
    try:
        # Find team where this student is a member
        # Since members is TextField, we search in it
        teams = Team.objects.all()
        student_team = None
        
        for team in teams:
            # Check if student name or ID is in members or member_names
            member_names = team.member_names or ''
            members = team.members or ''
            
            if (student_name in member_names) or (student_info.get('id') in members):
                student_team = team
                break
        
        if not student_team:
            return render(request, "student/student_result.html", {
                "allocation": None,
                "error": "You are not associated with any team.",
                "student": student_info
            })
        
        # Get allocation for this team
        allocation = AllocationResult.objects.filter(team_name=student_team.project_title).first()
        
        if not allocation:
            return render(request, "student/student_result.html", {
                "allocation": None,
                "error": "Mentor allocation not found for your team.",
                "team": student_team,
                "student": student_info
            })
        
        return render(request, "student/student_result.html", {
            "allocation": allocation,
            "team": student_team,
            "student": student_info
        })
        
    except Exception as e:
        return render(request, "student/student_result.html", {
            "allocation": None,
            "error": str(e),
            "student": student_info
        })


def mentor_result_view(request):
    """
    View for mentors to see their assigned teams.
    Uses mentor_name from session to match AllocationResult.
    """
    mentor_info = get_logged_in_mentor(request)
    
    # Check if mentor is logged in (by name)
    if not mentor_info['name']:
        # Fallback to Django user if available
        if request.user.is_authenticated:
            mentor_info['name'] = request.user.get_full_name() or request.user.username
        else:
            return redirect('/login/')
    
    try:
        # Search by mentor_name in AllocationResult
        search_name = mentor_info['name']
        
        # Get all allocations for this mentor
        allocations = AllocationResult.objects.filter(mentor_name=search_name)
        
        # If no allocations found, try case-insensitive search
        if not allocations.exists():
            allocations = AllocationResult.objects.filter(mentor_name__iexact=search_name)
        
        # If still no allocations, try contains search
        if not allocations.exists():
            allocations = AllocationResult.objects.filter(mentor_name__icontains=search_name.split()[0])  # Search by first name
        
        if not allocations.exists():
            return render(request, "mentor/mentor_result.html", {
                "allocations": [],
                "error": "No teams assigned to you yet.",
                "mentor": mentor_info,
                "unique_domains": [],
                "avg_similarity": 0,
                "active_count": 0
            })
        
        # Calculate statistics
        unique_domains = list(allocations.values_list('team_domain', flat=True).distinct())
        avg_similarity = allocations.aggregate(Avg('similarity_score'))['similarity_score__avg'] or 0
        
        context = {
            "allocations": allocations,
            "unique_domains": unique_domains,
            "avg_similarity": round(avg_similarity, 1),
            "active_count": allocations.count(),
            "mentor": mentor_info,
            "error": None
        }
        
        return render(request, "mentor/mentor_result.html", context)
        
    except Exception as e:
        print(f"Error in mentor_result_view: {str(e)}")
        return render(request, "mentor/mentor_result.html", {
            "allocations": [],
            "error": str(e),
            "mentor": mentor_info,
            "unique_domains": [],
            "avg_similarity": 0,
            "active_count": 0
        })


def zero_men(request):
    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")
    team_members = []
    team_name = None

    # Get the latest allocation for this mentor
    allocation = AllocationResult.objects.filter(mentor_name=mentor_name).order_by('-allocated_at').first()
    if allocation:
        team_name = allocation.team_name
        # Fetch the team object using project_title instead of team_name
        team = Team.objects.filter(project_title=team_name).first()
        if team and team.member_names:
            team_members = team.member_names.split(",")

    return render(request, "mentor/review_men/zero_men.html", {
        "mentor_name": mentor_name,
        "username": username,
        "team_name": team_name,
        "team_members": team_members,
    })

def one_men(request):
    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")
    team_members = []
    team_name = None

    # Get the latest allocation for this mentor
    allocation = AllocationResult.objects.filter(mentor_name=mentor_name).order_by('-allocated_at').first()
    if allocation:
        team_name = allocation.team_name
        # Fetch the team object using project_title instead of team_name
        team = Team.objects.filter(project_title=team_name).first()
        if team and team.member_names:
            team_members = team.member_names.split(",")

    return render(request, "mentor/review_men/1_men.html", {
        "mentor_name": mentor_name,
        "username": username,
        "team_name": team_name,
        "team_members": team_members,
    })

def two_men(request):
    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")
    team_members = []
    team_name = None

    # Get the latest allocation for this mentor
    allocation = AllocationResult.objects.filter(mentor_name=mentor_name).order_by('-allocated_at').first()
    if allocation:
        team_name = allocation.team_name

        # Fetch team using project_title instead of team_name
        team = Team.objects.filter(project_title=team_name).first()
        if team and team.member_names:
            team_members = team.member_names.split(",")

    return render(request, "mentor/review_men/2_men.html", {
        "mentor_name": mentor_name,
        "username": username,
        "team_name": team_name,
        "team_members": team_members,
    })


def three_men(request):
    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")
    team_members = []
    team_name = None

    # Get the latest allocation for this mentor
    allocation = AllocationResult.objects.filter(mentor_name=mentor_name).order_by('-allocated_at').first()
    if allocation:
        team_name = allocation.team_name

        # Fetch team using project_title instead of team_name
        team = Team.objects.filter(project_title=team_name).first()
        if team and team.member_names:
            team_members = team.member_names.split(",")

    return render(request, "mentor/review_men/3_men.html", {
        "mentor_name": mentor_name,
        "username": username,
        "team_name": team_name,
        "team_members": team_members,
    })



def serve_pdf(request, team_name, pdf_type):
    """
    Serve PDF from Cloudinary via direct URL (iframe-safe)
    """
    from allocation.models import TeamDocument  # nee create panna model (next step)

    doc = TeamDocument.objects.filter(
        team_name=team_name,
        doc_type=pdf_type
    ).first()

    if not doc or not doc.file_url:
        raise Http404("PDF not found")

    return redirect(doc.file_url)

import json
import os
from django.http import JsonResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.conf import settings

from .models import AllocationResult, ProjectFile, ProjectRemarks, ZerothReviewRemark, Annotation

import json
from django.http import JsonResponse
from django.shortcuts import render
from .models import ZerothReviewRemark, AllocationResult, ProjectFile

def zero_review(request):
    print("\n🟢 zero_review CALLED")
    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")
    print("mentor_name:", mentor_name)
    print("username:", username)
    print("method:", request.method)

    # ==================== POST: SAVE REMARKS ====================
    if request.method == "POST":
        try:
            raw_body = request.body.decode('utf-8')
            print("🔥 RAW REQUEST BODY (first 500 chars):", raw_body[:500])
            
            data = json.loads(raw_body)
            remarks = data.get("remarks", [])
            deleted = data.get("deleted", [])
            print("Incoming remarks count:", len(remarks))

            allocation = AllocationResult.objects.filter(mentor_name=mentor_name).first()
            if not allocation:
                return JsonResponse({"status": "fail", "message": "Team not found"}, status=404)

            team_name = allocation.team_name
            inserted = updated = deleted_count = 0

            # Handle deletions
            if deleted:
                for heading in deleted:
                    heading = heading.strip()
                    if not heading:
                        continue
                    count, _ = ZerothReviewRemark.objects.filter(
                        team_name=team_name, mentor_name=mentor_name, heading=heading
                    ).delete()
                    if count == 0:
                        count, _ = ZerothReviewRemark.objects.filter(
                            team_name=team_name, mentor_name=mentor_name, heading__icontains=heading[:50]
                        ).delete()
                    deleted_count += count

            # Handle saves/updates
            for r in remarks:
                heading = (r.get("heading") or "").strip()
                remark = (r.get("remark") or "").strip()
                color = r.get("color") or "#ffe066"
                
                # 🔥 CRITICAL: Extract coordinates
                coordinates = r.get("coordinates")
                print(f"\n🔥 PROCESSING REMARK: '{heading[:40]}...'")
                print(f"🔥 Raw coordinates from frontend: {type(coordinates)} = {coordinates is not None}")
                
                if not heading or not remark:
                    print("Skipping - missing heading or remark")
                    continue

                # Convert coordinates to JSON string
                if coordinates is None:
                    coords_json = "{}"
                elif isinstance(coordinates, dict):
                    coords_json = json.dumps(coordinates)
                    print(f"🔥 Dict converted to JSON string, length: {len(coords_json)}")
                elif isinstance(coordinates, str):
                    coords_json = coordinates
                else:
                    coords_json = "{}"

                print(f"🔥 coords_json to save (first 100 chars): {coords_json[:100]}")

                # 🔥 METHOD 1: Try update_or_create first
                try:
                    obj, created = ZerothReviewRemark.objects.update_or_create(
                        team_name=team_name,
                        mentor_name=mentor_name,
                        heading=heading,
                        defaults={
                            "remark": remark,
                            "color": color,
                            "coordinates": coords_json,
                            "file_type": "abstract"
                        },
                    )
                    print(f"🔥 update_or_create result: created={created}, id={obj.id}")
                    print(f"🔥 Saved coordinates in DB: {obj.coordinates[:100] if obj.coordinates else 'EMPTY'}")
                    
                except Exception as e:
                    print(f"🔥 update_or_create failed: {e}")
                    # 🔥 METHOD 2: Fallback - delete and recreate
                    ZerothReviewRemark.objects.filter(
                        team_name=team_name,
                        mentor_name=mentor_name,
                        heading=heading
                    ).delete()
                    
                    obj = ZerothReviewRemark.objects.create(
                        team_name=team_name,
                        mentor_name=mentor_name,
                        heading=heading,
                        remark=remark,
                        color=color,
                        coordinates=coords_json,
                        file_type="abstract"
                    )
                    print(f"🔥 Created new record after delete, id={obj.id}")
                    created = True

                if created:
                    inserted += 1
                else:
                    updated += 1

            return JsonResponse({
                "status": "success",
                "inserted": inserted,
                "updated": updated,
                "deleted": deleted_count,
            })

        except Exception as e:
            print("❌ POST ERROR:", e)
            import traceback
            traceback.print_exc()
            return JsonResponse({"status": "fail", "message": str(e)}, status=500)

    # ==================== GET: DISPLAY PAGE ====================
    allocation = AllocationResult.objects.filter(mentor_name=mentor_name).first()
    if not allocation:
        return render(request, "mentor/review_men/men_doc/zero_paper/zero_review.html")

    team_name = allocation.team_name
    print("Team:", team_name)

    # Load remarks
    saved_remarks = ZerothReviewRemark.objects.filter(
        team_name=team_name, mentor_name=mentor_name
    ).order_by("id")
    print("🔥 Loaded remarks from DB:", saved_remarks.count())

    # Parse coordinates
    for r in saved_remarks:
        raw_coords = r.coordinates
        print(f"\n🔥 DB Record: heading='{r.heading[:40]}...'")
        print(f"🔥 Raw coordinates from DB: {repr(raw_coords) if raw_coords else 'None/Empty'}")
        
        try:
            if raw_coords and raw_coords.strip() and raw_coords != "{}":
                r.parsed_coordinates = json.loads(raw_coords)
                print(f"🔥 Successfully parsed coordinates: {str(r.parsed_coordinates)[:100]}")
            else:
                r.parsed_coordinates = {}
                print(f"🔥 Empty or default coordinates, set to {{}}")
        except Exception as e:
            print(f"❌ Failed to parse coordinates: {e}")
            r.parsed_coordinates = {}

    # Get PDF URL
    project_file = ProjectFile.objects.filter(team_name=team_name, file_type="abstract").first()
    
    if not project_file:
        return render(request, "mentor/review_men/men_doc/zero_paper/zero_review.html", {
            "zero_review": False
        })

    pdf_url = project_file.cloudinary_url

    return render(
        request,
        "mentor/review_men/men_doc/zero_paper/zero_review.html",
        {
            "mentor_name": mentor_name,
            "username": username,
            "team_name": team_name,
            "pdf_url": pdf_url,
            "saved_remarks": saved_remarks,
            "zero_review": True,
        },
    )

# ─────────────────────────────────────────────
# NEW: API - Save Annotation
# ─────────────────────────────────────────────
@csrf_exempt
@require_http_methods(["POST"])
def save_annotation(request):
    """
    Save a single annotation (highlight or comment) to the database.
    Expects JSON body with annotation data.
    """
    try:
        data = json.loads(request.body)
        mentor_name = request.session.get("mentor_name")
        
        allocation = AllocationResult.objects.filter(mentor_name=mentor_name).first()
        if not allocation:
            return JsonResponse({"status": "fail", "message": "No allocation found"}, status=404)

        team_name = allocation.team_name
        project_file = ProjectFile.objects.filter(
            team_name=team_name,
            file_type=data.get("doc_type", "abstract")
        ).first()

        if not project_file:
            return JsonResponse({"status": "fail", "message": "Project file not found"}, status=404)

        # Create the annotation
        annotation = Annotation.objects.create(
            team=project_file,
            page_number=data.get("page", 1),
            annotation_type=data.get("annotation_type", "highlight"),
            x=float(data.get("x", 0)),
            y=float(data.get("y", 0)),
            width=float(data.get("width", 0)),
            height=float(data.get("height", 0)),
            color=data.get("color", "#FFFF00"),
            selected_text=data.get("selected_text", ""),
            comment=data.get("comment", ""),
            mentor=mentor_name
        )

        return JsonResponse({
            "status": "success",
            "annotation_id": annotation.id,
            "message": "Annotation saved"
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({"status": "fail", "message": str(e)}, status=500)


# ─────────────────────────────────────────────
# NEW: API - Get Annotations
# ─────────────────────────────────────────────
def get_annotations(request):
    """
    Get all annotations for a specific document and mentor.
    Query params: doc_type (abstract/pdf), team_name
    """
    mentor_name = request.session.get("mentor_name")
    doc_type = request.GET.get("doc_type", "abstract")
    team_name = request.GET.get("team_name")

    if not team_name:
        allocation = AllocationResult.objects.filter(mentor_name=mentor_name).first()
        if not allocation:
            return JsonResponse({"status": "fail", "message": "No allocation"}, status=404)
        team_name = allocation.team_name

    project_file = ProjectFile.objects.filter(
        team_name=team_name,
        file_type=doc_type
    ).first()

    if not project_file:
        return JsonResponse({"annotations": []})

    annotations = Annotation.objects.filter(
        team=project_file,
        mentor=mentor_name
    ).order_by('page_number', 'created_at')

    data = []
    for ann in annotations:
        data.append({
            "id": ann.id,
            "page": ann.page_number,
            "annotation_type": ann.annotation_type,
            "x": ann.x,
            "y": ann.y,
            "width": ann.width,
            "height": ann.height,
            "color": ann.color,
            "selected_text": ann.selected_text,
            "comment": ann.comment,
            "mentor": ann.mentor,
            "created_at": ann.created_at.isoformat()
        })

    return JsonResponse({
        "status": "success",
        "annotations": data,
        "count": len(data)
    })


# ─────────────────────────────────────────────
# NEW: API - Delete Annotation
# ─────────────────────────────────────────────
@csrf_exempt
@require_http_methods(["POST"])
def delete_annotation(request):
    """
    Delete a single annotation by ID.
    """
    try:
        data = json.loads(request.body)
        annotation_id = data.get("annotation_id")
        mentor_name = request.session.get("mentor_name")

        if not annotation_id:
            return JsonResponse({"status": "fail", "message": "annotation_id required"}, status=400)

        deleted, _ = Annotation.objects.filter(
            id=annotation_id,
            mentor=mentor_name
        ).delete()

        if deleted:
            return JsonResponse({"status": "success", "message": "Annotation deleted"})
        else:
            return JsonResponse({"status": "fail", "message": "Annotation not found"}, status=404)

    except Exception as e:
        return JsonResponse({"status": "fail", "message": str(e)}, status=500)


# ─────────────────────────────────────────────
# NEW: API - Update Annotation (Comment)
# ─────────────────────────────────────────────
@csrf_exempt
@require_http_methods(["POST"])
def update_annotation(request):
    """
    Update an annotation's comment or color.
    """
    try:
        data = json.loads(request.body)
        annotation_id = data.get("annotation_id")
        mentor_name = request.session.get("mentor_name")

        annotation = Annotation.objects.filter(
            id=annotation_id,
            mentor=mentor_name
        ).first()

        if not annotation:
            return JsonResponse({"status": "fail", "message": "Annotation not found"}, status=404)

        if "comment" in data:
            annotation.comment = data["comment"]
        if "color" in data:
            annotation.color = data["color"]

        annotation.save()

        return JsonResponse({
            "status": "success",
            "message": "Annotation updated"
        })

    except Exception as e:
        return JsonResponse({"status": "fail", "message": str(e)}, status=500)


# ─────────────────────────────────────────────
# OPTIONAL: Export Annotated PDF (Phase 6)
# ─────────────────────────────────────────────
def export_annotated_pdf(request):
    """
    Generate a PDF with annotations burned in.
    Called ONLY when user explicitly clicks 'Download Reviewed PDF'.
    """
    import requests
    from io import BytesIO
    
    mentor_name = request.session.get("mentor_name")
    doc_type = request.GET.get("doc_type", "abstract")
    
    allocation = AllocationResult.objects.filter(mentor_name=mentor_name).first()
    if not allocation:
        return JsonResponse({"status": "fail", "message": "No allocation"}, status=404)

    team_name = allocation.team_name
    
    project_file = ProjectFile.objects.filter(
        team_name=team_name,
        file_type=doc_type
    ).first()

    if not project_file:
        return JsonResponse({"status": "fail", "message": "File not found"}, status=404)

    # Get original PDF
    pdf_response = requests.get(project_file.cloudinary_url, timeout=30)
    if pdf_response.status_code != 200:
        return JsonResponse({"status": "fail", "message": "Could not fetch PDF"}, status=500)

    # Get annotations
    annotations = Annotation.objects.filter(
        team=project_file,
        mentor=mentor_name
    ).order_by('page_number')

    # TODO: Use reportlab or pypdf to overlay annotations on the PDF
    # Placeholder - implement actual PDF generation using reportlab, pypdf, or pdfrw
    
    return JsonResponse({
        "status": "success",
        "message": "Export functionality ready - implement PDF overlay logic",
        "annotations_count": annotations.count()
    })


# ─────────────────────────────────────────────
# KEEP: Rename your OLD zero_review to this for backup
# ─────────────────────────────────────────────
def zero_review_legacy(request):
    """
    BACKUP: Old HTML-conversion zero_review.
    Keep this during transition, remove once PDF.js version is stable.
    """
    # === PASTE YOUR ENTIRE OLD zero_review() CODE HERE ===
    pass

from django.http import FileResponse, Http404

def serve_temp_html(request, team, filename):
    html_path = os.path.join(
        settings.MEDIA_ROOT,
        "temp_html",
        team,
        filename
    )
    if not os.path.exists(html_path):
        raise Http404("HTML file not found")
    return FileResponse(
        open(html_path, "rb"),
        content_type="text/html"
    )

import json
import os
import requests
import subprocess
import re
from django.http import JsonResponse
from django.shortcuts import render
from django.conf import settings
from .models import ZerothReviewRemark, AllocationResult, ProjectFile

def zero_base(request):
    print("\n🟢 zero_base CALLED")
    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")
    print("mentor_name:", mentor_name)
    print("username:", username)
    print("method:", request.method)

    # ==================== POST: SAVE REMARKS ====================
    if request.method == "POST":
        try:
            raw_body = request.body.decode('utf-8')
            data = json.loads(raw_body)
            remarks = data.get("remarks", [])
            deleted = data.get("deleted", [])
            print("Incoming remarks:", len(remarks))
            print("Deleted headings:", len(deleted))

            allocation = AllocationResult.objects.filter(mentor_name=mentor_name).first()
            if not allocation:
                return JsonResponse({"status": "fail", "message": "Team not found"}, status=404)

            team_name = allocation.team_name
            inserted = updated = deleted_count = 0

            # Handle deletions
            if deleted:
                for heading in deleted:
                    heading = heading.strip()
                    if not heading:
                        continue
                    count, _ = ZerothReviewRemark.objects.filter(
                        team_name=team_name, mentor_name=mentor_name, 
                        heading=heading, file_type="pdf"
                    ).delete()
                    if count == 0:
                        count, _ = ZerothReviewRemark.objects.filter(
                            team_name=team_name, mentor_name=mentor_name,
                            heading__icontains=heading[:50], file_type="pdf"
                        ).delete()
                    deleted_count += count

            # Handle saves/updates
            for r in remarks:
                heading = (r.get("heading") or "").strip()
                remark = (r.get("remark") or "").strip()
                color = r.get("color") or "#ffe066"
                coordinates = r.get("coordinates") or {}
                
                print(f"🔥 Processing remark '{heading[:40]}...' with coordinates:", coordinates)
                
                if not heading or not remark:
                    continue

                # Convert coordinates to JSON string
                if isinstance(coordinates, dict):
                    coords_json = json.dumps(coordinates)
                elif isinstance(coordinates, str):
                    coords_json = coordinates
                else:
                    coords_json = "{}"

                obj, created = ZerothReviewRemark.objects.update_or_create(
                    team_name=team_name,
                    mentor_name=mentor_name,
                    heading=heading,
                    file_type="pdf",  # 🔥 BASE PAPER = pdf type
                    defaults={
                        "remark": remark,
                        "color": color,
                        "coordinates": coords_json,
                    },
                )
                print(f"✅ Saved record ID={obj.id}, coords={obj.coordinates[:100] if obj.coordinates else 'EMPTY'}")
                
                if created:
                    inserted += 1
                else:
                    updated += 1

            return JsonResponse({
                "status": "success",
                "inserted": inserted,
                "updated": updated,
                "deleted": deleted_count,
            })

        except Exception as e:
            print("❌ POST ERROR:", e)
            import traceback
            traceback.print_exc()
            return JsonResponse({"status": "fail", "message": str(e)}, status=500)

    # ==================== GET: DISPLAY PAGE ====================
    allocation = AllocationResult.objects.filter(mentor_name=mentor_name).first()
    if not allocation:
        return render(request, "mentor/review_men/men_doc/zero_paper/zero_base.html")

    team_name = allocation.team_name
    print("Team:", team_name)

    # Load remarks for BASE PAPER (file_type="pdf")
    saved_remarks = ZerothReviewRemark.objects.filter(
        team_name=team_name, mentor_name=mentor_name, file_type="pdf"
    ).order_by("id")
    print("🔥 Loaded base paper remarks:", saved_remarks.count())

    # Parse coordinates
    for r in saved_remarks:
        try:
            if r.coordinates and r.coordinates.strip() and r.coordinates != "{}":
                r.parsed_coordinates = json.loads(r.coordinates)
            else:
                r.parsed_coordinates = {}
        except Exception as e:
            print(f"❌ Failed to parse coordinates: {e}")
            r.parsed_coordinates = {}

    # Get PDF URL for base paper (report)
    project_file = ProjectFile.objects.filter(
        team_name=team_name, file_type="pdf"
    ).first()
    
    if not project_file:
        return render(request, "mentor/review_men/men_doc/zero_paper/zero_base.html", {
            "report_available": False
        })

    pdf_url = project_file.cloudinary_url

    return render(
        request,
        "mentor/review_men/men_doc/zero_paper/zero_base.html",
        {
            "mentor_name": mentor_name,
            "username": username,
            "team_name": team_name,
            "pdf_url": pdf_url,
            "saved_remarks": saved_remarks,
            "report_available": True,
        },
    )

def zero_ppt(request):
    print("\n🟢 zero_ppt CALLED")

    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")

    print("mentor_name:", mentor_name)
    print("username:", username)

    ppt_url = None
    team_name = None

    # =====================================================
    # 🔹 GET ALLOCATED TEAM (same as zero_review)
    # =====================================================
    allocation = AllocationResult.objects.filter(
        mentor_name=mentor_name
    ).first()

    if not allocation:
        print("❌ No team allocated")
        return render(
            request,
            "mentor/review_men/men_doc/zero_paper/zero_ppt.html",
            {
                "mentor_name": mentor_name,
                "username": username,
                "ppt_path": None,
                "team_name": None,
            }
        )

    team_name = allocation.team_name
    print("✔ Team:", team_name)

    # =====================================================
    # 🔥 FETCH PPT FROM ProjectFile (LIKE zero_review)
    # =====================================================
    project_file = ProjectFile.objects.filter(
        team_name=team_name,
        file_type="ppt"      # 👈 IMPORTANT
    ).first()

    if not project_file:
        print("❌ PPT not uploaded in ProjectFile")
        return render(
            request,
            "mentor/review_men/men_doc/zero_paper/zero_ppt.html",
            {
                "mentor_name": mentor_name,
                "username": username,
                "ppt_path": None,
                "team_name": team_name,
            }
        )

    ppt_url = project_file.cloudinary_url
    print("✔ PPT Cloudinary URL:", ppt_url)

    # =====================================================
    # FINAL RENDER
    # =====================================================
    return render(
        request,
        "mentor/review_men/men_doc/zero_paper/zero_ppt.html",
        {
            "mentor_name": mentor_name,
            "username": username,
            "ppt_path": ppt_url,   # 👈 SAME VARIABLE USED IN TEMPLATE
            "team_name": team_name,
        }
    )


def one_ppt(request):
    print("\n🟢 one_ppt CALLED")

    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")

    print("mentor_name:", mentor_name)
    print("username:", username)

    ppt_url = None
    team_name = None

    # =====================================================
    # 🔹 GET ALLOCATED TEAM (same as zero_ppt)
    # =====================================================
    allocation = AllocationResult.objects.filter(
        mentor_name=mentor_name
    ).first()

    if not allocation:
        print("❌ No team allocated")
        return render(
            request,
            "mentor/review_men/men_doc/first_paper/one_ppt.html",
            {
                "mentor_name": mentor_name,
                "username": username,
                "ppt_path": None,
                "team_name": None,
            }
        )

    team_name = allocation.team_name
    print("✔ Team:", team_name)

    # =====================================================
    # 🔥 FETCH PPT FROM ProjectFile (LIKE zero_ppt)
    # =====================================================
    project_file = ProjectFile.objects.filter(
        team_name=team_name,
        review_type="one",  # 👈 IMPORTANT: Review 1
        file_type="ppt"
    ).first()

    if not project_file:
        print("❌ PPT not uploaded in ProjectFile for Review 1")
        return render(
            request,
            "mentor/review_men/men_doc/first_paper/one_ppt.html",
            {
                "mentor_name": mentor_name,
                "username": username,
                "ppt_path": None,
                "team_name": team_name,
            }
        )

    ppt_url = project_file.cloudinary_url
    print("✔ PPT Cloudinary URL for Review 1:", ppt_url)

    # =====================================================
    # FINAL RENDER
    # =====================================================
    return render(
        request,
        "mentor/review_men/men_doc/first_paper/one_ppt.html",
        {
            "mentor_name": mentor_name,
            "username": username,
            "ppt_path": ppt_url,
            "team_name": team_name,
        }
    )


def two_ppt(request):
    print("\n🟢 two_ppt CALLED")

    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")

    print("mentor_name:", mentor_name)
    print("username:", username)

    ppt_url = None
    team_name = None

    # =====================================================
    # 🔹 GET ALLOCATED TEAM (same as zero_ppt)
    # =====================================================
    allocation = AllocationResult.objects.filter(
        mentor_name=mentor_name
    ).first()

    if not allocation:
        print("❌ No team allocated")
        return render(
            request,
            "mentor/review_men/men_doc/second_paper/two_ppt.html",
            {
                "mentor_name": mentor_name,
                "username": username,
                "ppt_path": None,
                "team_name": None,
            }
        )

    team_name = allocation.team_name
    print("✔ Team:", team_name)

    # =====================================================
    # 🔥 FETCH PPT FROM ProjectFile (LIKE zero_ppt)
    # =====================================================
    project_file = ProjectFile.objects.filter(
        team_name=team_name,
        review_type="two",  # 👈 IMPORTANT: Review 2
        file_type="ppt"
    ).first()

    if not project_file:
        print("❌ PPT not uploaded in ProjectFile for Review 2")
        return render(
            request,
            "mentor/review_men/men_doc/second_paper/two_ppt.html",
            {
                "mentor_name": mentor_name,
                "username": username,
                "ppt_path": None,
                "team_name": team_name,
            }
        )

    ppt_url = project_file.cloudinary_url
    print("✔ PPT Cloudinary URL for Review 2:", ppt_url)

    # =====================================================
    # FINAL RENDER
    # =====================================================
    return render(
        request,
        "mentor/review_men/men_doc/second_paper/two_ppt.html",
        {
            "mentor_name": mentor_name,
            "username": username,
            "ppt_path": ppt_url,
            "team_name": team_name,
        }
    )


def three_ppt(request):
    print("\n🟢 three_ppt CALLED")

    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")

    print("mentor_name:", mentor_name)
    print("username:", username)

    ppt_url = None
    team_name = None

    # =====================================================
    # 🔹 GET ALLOCATED TEAM (same as zero_ppt)
    # =====================================================
    allocation = AllocationResult.objects.filter(
        mentor_name=mentor_name
    ).first()

    if not allocation:
        print("❌ No team allocated")
        return render(
            request,
            "mentor/review_men/men_doc/third_paper/three_ppt.html",
            {
                "mentor_name": mentor_name,
                "username": username,
                "ppt_path": None,
                "team_name": None,
            }
        )

    team_name = allocation.team_name
    print("✔ Team:", team_name)

    # =====================================================
    # 🔥 FETCH PPT FROM ProjectFile (LIKE zero_ppt)
    # =====================================================
    project_file = ProjectFile.objects.filter(
        team_name=team_name,
        review_type="three",  # 👈 IMPORTANT: Review 3
        file_type="ppt"
    ).first()

    if not project_file:
        print("❌ PPT not uploaded in ProjectFile for Review 3")
        return render(
            request,
            "mentor/review_men/men_doc/third_paper/three_ppt.html",
            {
                "mentor_name": mentor_name,
                "username": username,
                "ppt_path": None,
                "team_name": team_name,
            }
        )

    ppt_url = project_file.cloudinary_url
    print("✔ PPT Cloudinary URL for Review 3:", ppt_url)

    # =====================================================
    # FINAL RENDER
    # =====================================================
    return render(
        request,
        "mentor/review_men/men_doc/third_paper/three_ppt.html",
        {
            "mentor_name": mentor_name,
            "username": username,
            "ppt_path": ppt_url,
            "team_name": team_name,
        }
    )

def zero_ma(request, team_name):
    team_members = []
    
    # Fetch team object using project_title
    team = Team.objects.filter(project_title=team_name).first()
    if team and team.member_names:
        # Convert comma-separated string to list
        team_members = team.member_names.split(",")

    # Render the zero_ma page for a specific team
    return render(request, 'mentor/review_men/men_ma/zero_ma.html', {
        'team_name': team_name,
        'team_members': team_members
    })

def one_ma(request, team_name):
    team_members = []
    
    # Fetch team object using project_title
    team = Team.objects.filter(project_title=team_name).first()
    if team and team.member_names:
        # Convert comma-separated string to list
        team_members = team.member_names.split(",")

    # Render the Review 1 (mentor assessment) page for a specific team
    return render(request, 'mentor/review_men/men_ma/one_ma.html', {
        'team_name': team_name,
        'team_members': team_members
    })


def two_ma(request, team_name):
    team_members = []
    
    # Fetch team object using project_title
    team = Team.objects.filter(project_title=team_name).first()
    if team and team.member_names:
        # Convert comma-separated string to list
        team_members = team.member_names.split(",")

    # Render the Review 1 (mentor assessment) page for a specific team
    return render(request, 'mentor/review_men/men_ma/two_ma.html', {
        'team_name': team_name,
        'team_members': team_members
    })

def three_ma(request, team_name):
    team_members = []
    
    # Fetch team object using project_title
    team = Team.objects.filter(project_title=team_name).first()
    if team and team.member_names:
        # Convert comma-separated string to list
        team_members = team.member_names.split(",")

    # Render the Review 3 mentor assessment page
    return render(request, 'mentor/review_men/men_ma/three_ma.html', {
        'team_name': team_name,
        'team_members': team_members
    })


def men_ppt(request):
    return render(request, "mentor/review_men/men_doc/first_paper/ppt.html")

def zero_stu(request):
    student_name = request.session.get("student_name")
    print(student_name)
    username = request.session.get("username")
    team = Team.objects.filter(member_names__icontains=student_name).first()
    if not team:
        return JsonResponse({"status": "fail", "message": "Team not found for this student"}, status=404)

    project_title = team.project_title.replace(" ", "_")  # sanitize for filename

    return render(request, "student/review/zero_stu.html", {
        "student_name": student_name,
        "username": username,
        "team_name": project_title,
    })


# ======================================================
# 🔹 REVIEW 1 — STUDENT
# ======================================================
def one_stu(request):
    student_name = request.session.get("student_name")
    print(student_name)
    username = request.session.get("username")
    
    if not student_name:
        return redirect("login")

    team = Team.objects.filter(member_names__icontains=student_name).first()
    if not team:
        return JsonResponse({"status": "fail", "message": "Team not found for this student"}, status=404)

    project_title = team.project_title.replace(" ", "_")  # sanitize for filename

    # 🔁 Check for existing PPT in Cloudinary only (no DB model)
    existing_ppt_url = None
    try:
        # Search in cloudinary folder for existing file
        search_result = cloudinary.Search().expression(
            f"folder:review1_ppt AND filename:{project_title}*"
        ).execute()
        
        if search_result.get("resources"):
            existing_ppt_url = search_result["resources"][0]["secure_url"]
    except Exception as e:
        print(f"Cloudinary search error: {e}")

    if request.method == "POST":
        ppt_file = request.FILES.get("pptFile")

        if not ppt_file:
            return JsonResponse({"status": "fail", "message": "No PPT uploaded"})

        # Upload to Cloudinary directly (no DB storage)
        upload = cloudinary.uploader.upload(
            ppt_file,
            resource_type="raw",
            folder="review1_ppt",
            public_id=f"{project_title}_ppt_{int(time.time())}"  # unique name with timestamp
        )

        return JsonResponse({
            "status": "success",
            "ppt_url": upload["secure_url"]
        })

    # Simple render like zero_stu - no complex DB queries
    return render(request, "student/review/1_stu.html", {
        "student_name": student_name,
        "username": username,
        "team_name": project_title,
        "ppt_url": existing_ppt_url,  # Pass URL directly, not DB object
    })

# ======================================================
# 🔹 REVIEW 2 — STUDENT
# ======================================================
def two_stu(request):
    student_name = request.session.get("student_name")
    print(student_name)
    username = request.session.get("username")
    
    if not student_name:
        return redirect("login")

    team = Team.objects.filter(member_names__icontains=student_name).first()
    if not team:
        return JsonResponse({"status": "fail", "message": "Team not found for this student"}, status=404)

    project_title = team.project_title.replace(" ", "_")  # sanitize for filename

    # 🔁 Check for existing PPT in Cloudinary only (no DB model)
    existing_ppt_url = None
    try:
        # Search in cloudinary folder for existing file
        search_result = cloudinary.Search().expression(
            f"folder:review2_ppt AND filename:{project_title}*"
        ).execute()
        
        if search_result.get("resources"):
            existing_ppt_url = search_result["resources"][0]["secure_url"]
    except Exception as e:
        print(f"Cloudinary search error: {e}")

    if request.method == "POST":
        ppt_file = request.FILES.get("pptFile")

        if not ppt_file:
            return JsonResponse({"status": "fail", "message": "No PPT uploaded"})

        # Upload to Cloudinary directly (no DB storage)
        upload = cloudinary.uploader.upload(
            ppt_file,
            resource_type="raw",
            folder="review2_ppt",
            public_id=f"{project_title}_ppt_{int(time.time())}"  # unique name with timestamp
        )

        return JsonResponse({
            "status": "success",
            "ppt_url": upload["secure_url"]
        })

    # Simple render like zero_stu - no complex DB queries
    return render(request, "student/review/2_stu.html", {
        "student_name": student_name,
        "username": username,
        "team_name": project_title,
        "ppt_url": existing_ppt_url,  # Pass URL directly, not DB object
    })

# ======================================================
# 🔹 REVIEW 3 — STUDENT
# ======================================================
def three_stu(request):
    student_name = request.session.get("student_name")
    print(student_name)
    username = request.session.get("username")
    
    if not student_name:
        return redirect("login")

    team = Team.objects.filter(member_names__icontains=student_name).first()
    if not team:
        return JsonResponse({"status": "fail", "message": "Team not found for this student"}, status=404)

    project_title = team.project_title.replace(" ", "_")  # sanitize for filename

    # 🔁 Check for existing PPT in Cloudinary only (no DB model)
    existing_ppt_url = None
    try:
        # Search in cloudinary folder for existing file
        search_result = cloudinary.Search().expression(
            f"folder:review3_ppt AND filename:{project_title}*"
        ).execute()
        
        if search_result.get("resources"):
            existing_ppt_url = search_result["resources"][0]["secure_url"]
    except Exception as e:
        print(f"Cloudinary search error: {e}")

    if request.method == "POST":
        ppt_file = request.FILES.get("pptFile")

        if not ppt_file:
            return JsonResponse({"status": "fail", "message": "No PPT uploaded"})

        # Upload to Cloudinary directly (no DB storage)
        upload = cloudinary.uploader.upload(
            ppt_file,
            resource_type="raw",
            folder="review3_ppt",
            public_id=f"{project_title}_ppt_{int(time.time())}"  # unique name with timestamp
        )

        return JsonResponse({
            "status": "success",
            "ppt_url": upload["secure_url"]
        })

    # Simple render like zero_stu - no complex DB queries
    return render(request, "student/review/3_stu.html", {
        "student_name": student_name,
        "username": username,
        "team_name": project_title,
        "ppt_url": existing_ppt_url,  # Pass URL directly, not DB object
    })

def mentor_list(request):
    mentors = Mentor.objects.all()
    return render(request, "coordinator/men_list.html", {"mentors": mentors})

from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
from django.shortcuts import render, get_object_or_404
from django.contrib import messages
import json

@require_http_methods(["GET", "POST"])
def team_list(request):
    if request.method == "POST":
        approved_ids = request.POST.getlist('approved_ids')
        modified_ids = request.POST.getlist('modified_ids')
        
        approved_count = 0
        modified_count = 0
        
        # Approve teams
        for team_id in approved_ids:
            team = Team.objects.get(id=team_id)
            team.status = 'approved'
            team.needs_update_problem = False
            team.needs_update_domain = False
            team.needs_update_members = False
            team.modification_reason = ''
            team.save()
            approved_count += 1
        
        # Handle modifications - clear only selected fields, save ticks and reason
        for team_id in modified_ids:
            team = Team.objects.get(id=team_id)
            fields_json = request.POST.get(f'modify_fields_{team_id}', '[]')
            fields = json.loads(fields_json)
            reason = request.POST.get(f'modify_reason_{team_id}', '')
            
            # Reset all flags first
            team.needs_update_problem = False
            team.needs_update_domain = False
            team.needs_update_members = False
            
            # Set only selected fields
            if 'problem_statement' in fields:
                team.project_title = ''
                team.needs_update_problem = True
            
            if 'domain' in fields:
                team.domain = ''
                team.needs_update_domain = True
            
            if 'team_members' in fields:
                leader = team.members.split(',')[0] if team.members else ''
                team.members = leader
                team.member_names = ''
                team.needs_update_members = True
            
            team.modification_reason = reason
            team.status = 'pending_update'
            team.save()
            modified_count += 1
        
        return JsonResponse({
            'success': True,
            'approved_count': approved_count,
            'modified_count': modified_count
        })
    
    teams = Team.objects.all().order_by('-created_at')
    
    # Count stats
    approved_count = teams.filter(status='approved').count()
    pending_count = teams.filter(status='pending').count()
    modify_count = teams.filter(status='pending_update').count()
    
    return render(request, 'coordinator/team_list.html', {
        'teams': teams,
        'approved_count': approved_count,
        'pending_count': pending_count,
        'modify_count': modify_count,
    })


def student_update_team(request):
    """
    Call this from student side after they update their team.
    Changes status from 'pending_update' to 'updated' to notify coordinator.
    """
    if request.method == "POST":
        try:
            body = request.body.decode("utf-8")
            data = json.loads(body) if body else {}
            team_id = data.get('team_id')
            
            team = get_object_or_404(Team, id=team_id)
            
            # Check if all requested updates are done
            all_updated = True
            if team.needs_update_problem and not team.project_title:
                all_updated = False
            if team.needs_update_domain and not team.domain:
                all_updated = False
            if team.needs_update_members and len(team.members.split(',')) < 2:
                all_updated = False
            
            if all_updated:
                team.status = 'updated'  # Notify coordinator that student has updated
                team.save()
                return JsonResponse({
                    'status': 'success',
                    'message': 'Team updated. Coordinator will be notified.'
                })
            
            return JsonResponse({
                'status': 'error',
                'message': 'Please complete all requested updates.'
            })
            
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})
    
    return JsonResponse({'status': 'error', 'message': 'Invalid request'})


def approve_team(request, team_id):
    """Direct approve from popup or other views"""
    team = get_object_or_404(Team, id=team_id)
    team.status = 'approved'
    team.needs_update_problem = False
    team.needs_update_domain = False
    team.needs_update_members = False
    team.modification_reason = ''
    team.save()
    
    messages.success(request, f"'{team.project_title}' approved successfully!")
    return redirect("team_list")

def modify_team(request, project_title):
    if request.method == "POST":
        change_type = request.POST.get("change_type")
        print("Modify triggered:", request.POST)

        team = get_object_or_404(Team, project_title=project_title)

        # 🧹 Step 1: Remove from Approved list if exists
        ApprovedTeam.objects.filter(project_title=project_title).delete()

        # 🧹 Step 2: Remove old modify request (avoid duplicates)
        ModifyRequest.objects.filter(project_title=project_title).delete()

        # 🧩 Step 3: Add new modify request
        ModifyRequest.objects.create(
            project_title=team.project_title,
            student_class=team.student_class,
            domain=team.domain,
            members=team.members,
            member_names=team.member_names,
            change_type=change_type
        )

        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return JsonResponse({"status": "success"})

        messages.success(request, f"Modification request for '{team.project_title}' ({change_type}) added successfully!")
        return redirect("team_list")



import io
import json
import cloudinary.uploader

from django.shortcuts import render, redirect
from django.http import JsonResponse
from django.db import transaction

from allocation.models import ZerothReviewRemark
from allocation.models import ProjectFile, Team


# ============================================
# 🔹 Helper Function: Upload to Cloudinary
# ============================================
from django.shortcuts import render, redirect
from django.http import JsonResponse
from allocation.models import Team, ProjectFile, ZerothReviewRemark,FirstReviewRemark,SecondReviewRemark,ThirdReviewRemark
import cloudinary
import json

# -----------------------------
# Helper: Upload to Cloudinary
# -----------------------------
def upload_to_cloudinary(file_obj, file_type, folder_name):
    try:
        print(f"DEBUG: Uploading {file_type} to Cloudinary...")
        result = cloudinary.uploader.upload(
            file_obj,
            resource_type="auto",       # Supports PDF, PPT, etc.

            folder=f"project_portal/Upload_docs/Zero_Review/{folder_name}",
            public_id=f"{folder_name}_{file_type}",
            overwrite=True,
            use_filename=True,
            unique_filename=False,
            access_mode="public"        # Ensure public access
        )
        file_url = result.get("secure_url")
        print(f"DEBUG: Uploaded {file_type} URL → {file_url}")
        return file_url
    except Exception as e:
        print(f"❌ Cloudinary upload failed for {file_type}: {e}")
        return None

def upload_to_cloudinary1(file_obj, file_type, folder_name):
    try:
        print(f"DEBUG: Uploading {file_type} to Cloudinary for Review 1...")
        result = cloudinary.uploader.upload(
            file_obj,
            resource_type="auto",       # Supports PDF, PPT, etc.

            folder=f"project_portal/Upload_docs/First_Review/{folder_name}",  # Changed to Review_1
            public_id=f"{folder_name}_{file_type}_review1",  # Added review1 identifier
            overwrite=True,
            use_filename=True,
            unique_filename=False,
            access_mode="public"        # Ensure public access
        )
        file_url = result.get("secure_url")
        print(f"DEBUG: Uploaded {file_type} URL for Review 1 → {file_url}")
        return file_url
    except Exception as e:
        print(f"❌ Cloudinary upload failed for {file_type} in Review 1: {e}")
        return None

def upload_to_cloudinary2(file_obj, file_type, folder_name):
    try:
        print(f"DEBUG: Uploading {file_type} to Cloudinary for Review 2...")
        result = cloudinary.uploader.upload(
            file_obj,
            resource_type="auto",       # Supports PDF, PPT, etc.

            folder=f"project_portal/Upload_docs/Second_Review/{folder_name}",  # Changed to Second_Review
            public_id=f"{folder_name}_{file_type}_review2",  # Added review2 identifier
            overwrite=True,
            use_filename=True,
            unique_filename=False,
            access_mode="public"        # Ensure public access
        )
        file_url = result.get("secure_url")
        print(f"DEBUG: Uploaded {file_type} URL for Review 2 → {file_url}")
        return file_url
    except Exception as e:
        print(f"❌ Cloudinary upload failed for {file_type} in Review 2: {e}")
        return None

def upload_to_cloudinary3(file_obj, file_type, folder_name):
    try:
        print(f"DEBUG: Uploading {file_type} to Cloudinary for Review 3...")
        result = cloudinary.uploader.upload(
            file_obj,
            resource_type="auto",       # Supports PDF, PPT, etc.

            folder=f"project_portal/Upload_docs/Third_Review/{folder_name}",  # Changed to Third_Review
            public_id=f"{folder_name}_{file_type}_review3",  # Added review3 identifier
            overwrite=True,
            use_filename=True,
            unique_filename=False,
            access_mode="public"        # Ensure public access
        )
        file_url = result.get("secure_url")
        print(f"DEBUG: Uploaded {file_type} URL for Review 3 → {file_url}")
        return file_url
    except Exception as e:
        print(f"❌ Cloudinary upload failed for {file_type} in Review 3: {e}")
        return None
# -----------------------------
# View: Student Upload (Zero Review)
# -----------------------------
import json
from django.shortcuts import render, redirect
from .models import Team, ProjectFile, ZerothReviewRemark, ProjectRemarks

def zero_ma1(request):
    # ---------------------------
    # 1️⃣ Get Student Session
    # ---------------------------
    student_name = request.session.get("student_name")
    username = request.session.get("username")
    print("DEBUG: Student session →", student_name, username)

    if not student_name:
        return redirect("login")

    # ---------------------------
    # 2️⃣ Find Student Team
    # ---------------------------
    team = Team.objects.filter(member_names__icontains=student_name).first()
    if not team:
        return render(request, "student/review/zero_ma.html", {
            "student_name": student_name,
            "username": username,
            "error": "Team not found"
        })

    team_title = team.project_title
    folder_name = team_title.replace(" ", "_")
    print("DEBUG: Found team →", team_title)

    # ---------------------------
    # 3️⃣ Handle POST → Upload Files
    # ---------------------------
    if request.method == "POST":
        # 🔥 Check if re-upload is allowed (mentor requested)
        reupload_allowed = team.reupload_allowed if team else False
        
        ppt_file = request.FILES.get("pptFile")
        pdf_file = request.FILES.get("pdfFile")
        abstract_file = request.FILES.get("abstractFile")
        print("DEBUG: Files received →", ppt_file, pdf_file, abstract_file)

        uploaded = {}

        # 🔥 If re-upload allowed, clear existing files first
        if reupload_allowed:
            ProjectFile.objects.filter(team_name=team_title, review_type="zero").delete()
            team.reupload_allowed = False  # Reset after re-upload
            team.save()
            print("DEBUG: Cleared existing files for re-upload")

        if ppt_file:
            uploaded["ppt"] = upload_to_cloudinary(ppt_file, "PPT", folder_name)

        if pdf_file:
            uploaded["pdf"] = upload_to_cloudinary(pdf_file, "Report", folder_name)

        if abstract_file:
            uploaded["abstract"] = upload_to_cloudinary(abstract_file, "Abstract", folder_name)

        print("DEBUG: Uploaded files dict →", uploaded)

        # ---------------------------
        # 4️⃣ Update ProjectFile Table
        # ---------------------------
        for ftype, url in uploaded.items():
            if url:
                obj, created = ProjectFile.objects.update_or_create(
                    team_name=team_title,
                    review_type="zero",
                    file_type=ftype,
                    defaults={"cloudinary_url": url}
                )
                print(f"DEBUG: ProjectFile {'created' if created else 'updated'} → {ftype}: {url}")

        return JsonResponse({
            "status": "success",
            "message": "Files uploaded to Cloudinary",
            "files": uploaded,
            "reupload": reupload_allowed
        })

    # ---------------------------
    # 5️⃣ GET → Fetch already uploaded files
    # ---------------------------
    uploaded_files = {}
    files_qs = ProjectFile.objects.filter(team_name=team_title, review_type="zero")
    for f in files_qs:
        uploaded_files[f.file_type] = f.cloudinary_url
    print("DEBUG: Uploaded files fetched →", uploaded_files)

    # ---------------------------
    # 6️⃣ Get Zeroth Review Remarks (GROUPED BY FILE TYPE)
    # ---------------------------
    remarks_by_type = {
        "abstract": [],
        "pdf": [],
        "ppt": []
    }
    
    remarks_qs = ZerothReviewRemark.objects.filter(team_name=team_title).order_by("created_at")
    
    for r in remarks_qs:
        # 🔥 PARSE COORDINATES for frontend rendering
        parsed_coords = {}
        try:
            if r.coordinates and r.coordinates.strip() and r.coordinates != "{}":
                if isinstance(r.coordinates, str):
                    parsed_coords = json.loads(r.coordinates)
                elif isinstance(r.coordinates, dict):
                    parsed_coords = r.coordinates
        except Exception as e:
            print(f"❌ Failed to parse coordinates: {e}")
            parsed_coords = {}
        
        remark_data = {
            "heading": r.heading,
            "remark": r.remark,
            "color": r.color,
            "created_at": r.created_at,
            "mentor_name": r.mentor_name,
            "coordinates": parsed_coords,  # 🔥 Pass parsed coordinates to template
        }
        
        # Determine file type
        file_type = getattr(r, 'file_type', 'abstract')
        if file_type in remarks_by_type:
            remarks_by_type[file_type].append(remark_data)
        else:
            remarks_by_type["abstract"].append(remark_data)
    
    print("DEBUG: Remarks fetched →", {k: len(v) for k, v in remarks_by_type.items()})

    # ---------------------------
    # 7️⃣ Get Highlighted PDFs from ProjectRemarks
    # ---------------------------
    highlighted_pdfs = {}
    
    remarks_files_qs = ProjectRemarks.objects.filter(
        team_name=team_title,
        review_type="zero"
    ).select_related('original_file')
    
    for remark_file in remarks_files_qs:
        file_type = remark_file.file_type
        highlighted_pdfs[file_type] = {
            "url": remark_file.cloudinary_url,
            "mentor_name": remark_file.mentor_name,
            "updated_at": remark_file.updated_at if hasattr(remark_file, 'updated_at') else None,
            "original_file_type": file_type
        }
    
    print("DEBUG: Highlighted PDFs fetched →", highlighted_pdfs)

    # ---------------------------
    # 8️⃣ Final Render
    # ---------------------------
    reupload_info = {
        "allowed": team.reupload_allowed if team else False,
        "reason": team.reupload_reason if team else "",
        "requested_by": team.reupload_requested_by if team else "",
        "requested_at": team.reupload_requested_at.strftime("%Y-%m-%d %H:%M") if team and team.reupload_requested_at else ""
    }
    
    return render(request, "student/review/zero_ma.html", {
        "student_name": student_name,
        "username": username,
        "team_name": team_title,
        "uploaded_files": uploaded_files,
        "remarks_by_type": remarks_by_type,
        "highlighted_pdfs": highlighted_pdfs,
        "reupload_info": reupload_info,  # 🔥 Pass re-upload info
    })
# ============================================
# 🔹 Student Zero Review File Upload View
# ============================================
def one_ma1(request):
    # ---------------------------
    # 1️⃣ Get Student Session
    # ---------------------------
    student_name = request.session.get("student_name")
    username = request.session.get("username")
    print("DEBUG: Student session →", student_name, username)

    if not student_name:
        return redirect("login")

    # ---------------------------
    # 2️⃣ Find Student Team
    # ---------------------------
    team = Team.objects.filter(member_names__icontains=student_name).first()
    if not team:
        return render(request, "student/review/1_ma.html", {
            "student_name": student_name,
            "username": username,
            "error": "Team not found"
        })

    team_title = team.project_title
    folder_name = team_title.replace(" ", "_")
    print("DEBUG: Found team →", team_title)

    # ---------------------------
    # 3️⃣ Handle POST → Upload PPT Only
    # ---------------------------
    if request.method == "POST":
        ppt_file = request.FILES.get("pptFile")
        print("DEBUG: File received →", ppt_file)

        uploaded = {}

        if ppt_file:
            uploaded["ppt"] = upload_to_cloudinary1(ppt_file, "PPT", folder_name)

        print("DEBUG: Uploaded file →", uploaded)

        # ---------------------------
        # 4️⃣ Update ProjectFile Table (Review 1)
        # ---------------------------
        for ftype, url in uploaded.items():
            if url:
                obj, created = ProjectFile.objects.update_or_create(
                    team_name=team_title,
                    review_type="one",  # Changed from "zero" to "one"
                    file_type=ftype,
                    defaults={"cloudinary_url": url}
                )
                print(f"DEBUG: ProjectFile {'created' if created else 'updated'} → {ftype}: {url}")

        return JsonResponse({
            "status": "success",
            "message": "PPT uploaded to Cloudinary",
            "files": uploaded
        })

    # ---------------------------
    # 5️⃣ GET → Fetch already uploaded PPT
    # ---------------------------
    uploaded_files = {}
    files_qs = ProjectFile.objects.filter(team_name=team_title, review_type="one")  # Changed to "one"
    for f in files_qs:
        uploaded_files[f.file_type] = f.cloudinary_url
    print("DEBUG: Uploaded files fetched →", uploaded_files)

    # ---------------------------
    # 6️⃣ Get First Review Remarks (GROUPED BY FILE TYPE)
    # ---------------------------
    remarks_by_type = {
        "ppt": []  # Only PPT for review 1
    }
    
    # Fetch all remarks for this team (Review 1)
    remarks_qs = FirstReviewRemark.objects.filter(team_name=team_title).order_by("created_at")  # Use FirstReviewRemark model
    
    for r in remarks_qs:
        remark_data = {
            "heading": r.heading,
            "remark": r.remark,
            "color": r.color,
            "created_at": r.created_at,
            "mentor_name": r.mentor_name if hasattr(r, 'mentor_name') else None
        }
        
        # Determine file type
        if hasattr(r, 'file_type') and r.file_type:
            if r.file_type in remarks_by_type:
                remarks_by_type[r.file_type].append(remark_data)
            else:
                remarks_by_type["ppt"].append(remark_data)
        else:
            # Default to PPT for review 1
            remarks_by_type["ppt"].append(remark_data)
    
    print("DEBUG: Remarks fetched →", {k: len(v) for k, v in remarks_by_type.items()})

    # ---------------------------
    # 7️⃣ Get Highlighted PPTs from ProjectRemarks (Review 1)
    # ---------------------------
    highlighted_pdfs = {}
    
    remarks_files_qs = ProjectRemarks.objects.filter(
        team_name=team_title,
        review_type="one"  # Changed to "one"
    ).select_related('original_file')
    
    for remark_file in remarks_files_qs:
        file_type = remark_file.file_type
        highlighted_pdfs[file_type] = {
            "url": remark_file.cloudinary_url,
            "mentor_name": remark_file.mentor_name,
            "updated_at": remark_file.updated_at if hasattr(remark_file, 'updated_at') else None,
            "original_file_type": file_type
        }
    
    print("DEBUG: Highlighted files fetched →", highlighted_pdfs)

    # ---------------------------
    # 8️⃣ Final Render
    # ---------------------------
    return render(request, "student/review/1_ma.html", {
        "student_name": student_name,
        "username": username,
        "team_name": team_title,
        "uploaded_files": uploaded_files,
        "remarks_by_type": remarks_by_type,
        "highlighted_pdfs": highlighted_pdfs,
    })

def two_ma1(request):
    # ---------------------------
    # 1️⃣ Get Student Session
    # ---------------------------
    student_name = request.session.get("student_name")
    username = request.session.get("username")
    print("DEBUG: Student session →", student_name, username)

    if not student_name:
        return redirect("login")

    # ---------------------------
    # 2️⃣ Find Student Team
    # ---------------------------
    team = Team.objects.filter(member_names__icontains=student_name).first()
    if not team:
        return render(request, "student/review/2_ma.html", {
            "student_name": student_name,
            "username": username,
            "error": "Team not found"
        })

    team_title = team.project_title
    folder_name = team_title.replace(" ", "_")
    print("DEBUG: Found team →", team_title)

    # ---------------------------
    # 3️⃣ Handle POST → Upload PPT Only
    # ---------------------------
    if request.method == "POST":
        ppt_file = request.FILES.get("pptFile")
        print("DEBUG: File received →", ppt_file)

        uploaded = {}

        if ppt_file:
            uploaded["ppt"] = upload_to_cloudinary2(ppt_file, "PPT", folder_name)

        print("DEBUG: Uploaded file →", uploaded)

        # ---------------------------
        # 4️⃣ Update ProjectFile Table (Review 2)
        # ---------------------------
        for ftype, url in uploaded.items():
            if url:
                obj, created = ProjectFile.objects.update_or_create(
                    team_name=team_title,
                    review_type="two",  # Changed to "two"
                    file_type=ftype,
                    defaults={"cloudinary_url": url}
                )
                print(f"DEBUG: ProjectFile {'created' if created else 'updated'} → {ftype}: {url}")

        return JsonResponse({
            "status": "success",
            "message": "PPT uploaded to Cloudinary",
            "files": uploaded
        })

    # ---------------------------
    # 5️⃣ GET → Fetch already uploaded PPT
    # ---------------------------
    uploaded_files = {}
    files_qs = ProjectFile.objects.filter(team_name=team_title, review_type="two")  # Changed to "two"
    for f in files_qs:
        uploaded_files[f.file_type] = f.cloudinary_url
    print("DEBUG: Uploaded files fetched →", uploaded_files)

    # ---------------------------
    # 6️⃣ Get Second Review Remarks (GROUPED BY FILE TYPE)
    # ---------------------------
    remarks_by_type = {
        "ppt": []  # Only PPT for review 2
    }
    
    # Fetch all remarks for this team (Review 2)
    remarks_qs = SecondReviewRemark.objects.filter(team_name=team_title).order_by("created_at")  # Use SecondReviewRemark model
    
    for r in remarks_qs:
        remark_data = {
            "heading": r.heading,
            "remark": r.remark,
            "color": r.color,
            "created_at": r.created_at,
            "mentor_name": r.mentor_name if hasattr(r, 'mentor_name') else None
        }
        
        # Determine file type
        if hasattr(r, 'file_type') and r.file_type:
            if r.file_type in remarks_by_type:
                remarks_by_type[r.file_type].append(remark_data)
            else:
                remarks_by_type["ppt"].append(remark_data)
        else:
            # Default to PPT for review 2
            remarks_by_type["ppt"].append(remark_data)
    
    print("DEBUG: Remarks fetched →", {k: len(v) for k, v in remarks_by_type.items()})

    # ---------------------------
    # 7️⃣ Get Highlighted PPTs from ProjectRemarks (Review 2)
    # ---------------------------
    highlighted_pdfs = {}
    
    remarks_files_qs = ProjectRemarks.objects.filter(
        team_name=team_title,
        review_type="two"  # Changed to "two"
    ).select_related('original_file')
    
    for remark_file in remarks_files_qs:
        file_type = remark_file.file_type
        highlighted_pdfs[file_type] = {
            "url": remark_file.cloudinary_url,
            "mentor_name": remark_file.mentor_name,
            "updated_at": remark_file.updated_at if hasattr(remark_file, 'updated_at') else None,
            "original_file_type": file_type
        }
    
    print("DEBUG: Highlighted files fetched →", highlighted_pdfs)

    # ---------------------------
    # 8️⃣ Final Render
    # ---------------------------
    return render(request, "student/review/2_ma.html", {
        "student_name": student_name,
        "username": username,
        "team_name": team_title,
        "uploaded_files": uploaded_files,
        "remarks_by_type": remarks_by_type,
        "highlighted_pdfs": highlighted_pdfs,
    })

def three_ma1(request):
    # ---------------------------
    # 1️⃣ Get Student Session
    # ---------------------------
    student_name = request.session.get("student_name")
    username = request.session.get("username")
    print("DEBUG: Student session →", student_name, username)

    if not student_name:
        return redirect("login")

    # ---------------------------
    # 2️⃣ Find Student Team
    # ---------------------------
    team = Team.objects.filter(member_names__icontains=student_name).first()
    if not team:
        return render(request, "student/review/3_ma.html", {
            "student_name": student_name,
            "username": username,
            "error": "Team not found"
        })

    team_title = team.project_title
    folder_name = team_title.replace(" ", "_")
    print("DEBUG: Found team →", team_title)

    # ---------------------------
    # 3️⃣ Handle POST → Upload PPT Only
    # ---------------------------
    if request.method == "POST":
        ppt_file = request.FILES.get("pptFile")
        print("DEBUG: File received →", ppt_file)

        uploaded = {}

        if ppt_file:
            uploaded["ppt"] = upload_to_cloudinary3(ppt_file, "PPT", folder_name)

        print("DEBUG: Uploaded file →", uploaded)

        # ---------------------------
        # 4️⃣ Update ProjectFile Table (Review 3)
        # ---------------------------
        for ftype, url in uploaded.items():
            if url:
                obj, created = ProjectFile.objects.update_or_create(
                    team_name=team_title,
                    review_type="three",  # Changed to "three"
                    file_type=ftype,
                    defaults={"cloudinary_url": url}
                )
                print(f"DEBUG: ProjectFile {'created' if created else 'updated'} → {ftype}: {url}")

        return JsonResponse({
            "status": "success",
            "message": "PPT uploaded to Cloudinary",
            "files": uploaded
        })

    # ---------------------------
    # 5️⃣ GET → Fetch already uploaded PPT
    # ---------------------------
    uploaded_files = {}
    files_qs = ProjectFile.objects.filter(team_name=team_title, review_type="three")  # Changed to "three"
    for f in files_qs:
        uploaded_files[f.file_type] = f.cloudinary_url
    print("DEBUG: Uploaded files fetched →", uploaded_files)

    # ---------------------------
    # 6️⃣ Get Third Review Remarks (GROUPED BY FILE TYPE)
    # ---------------------------
    remarks_by_type = {
        "ppt": []  # Only PPT for review 3
    }
    
    # Fetch all remarks for this team (Review 3)
    remarks_qs = ThirdReviewRemark.objects.filter(team_name=team_title).order_by("created_at")  # Use ThirdReviewRemark model
    
    for r in remarks_qs:
        remark_data = {
            "heading": r.heading,
            "remark": r.remark,
            "color": r.color,
            "created_at": r.created_at,
            "mentor_name": r.mentor_name if hasattr(r, 'mentor_name') else None
        }
        
        # Determine file type
        if hasattr(r, 'file_type') and r.file_type:
            if r.file_type in remarks_by_type:
                remarks_by_type[r.file_type].append(remark_data)
            else:
                remarks_by_type["ppt"].append(remark_data)
        else:
            # Default to PPT for review 3
            remarks_by_type["ppt"].append(remark_data)
    
    print("DEBUG: Remarks fetched →", {k: len(v) for k, v in remarks_by_type.items()})

    # ---------------------------
    # 7️⃣ Get Highlighted PPTs from ProjectRemarks (Review 3)
    # ---------------------------
    highlighted_pdfs = {}
    
    remarks_files_qs = ProjectRemarks.objects.filter(
        team_name=team_title,
        review_type="three"  # Changed to "three"
    ).select_related('original_file')
    
    for remark_file in remarks_files_qs:
        file_type = remark_file.file_type
        highlighted_pdfs[file_type] = {
            "url": remark_file.cloudinary_url,
            "mentor_name": remark_file.mentor_name,
            "updated_at": remark_file.updated_at if hasattr(remark_file, 'updated_at') else None,
            "original_file_type": file_type
        }
    
    print("DEBUG: Highlighted files fetched →", highlighted_pdfs)

    # ---------------------------
    # 8️⃣ Final Render
    # ---------------------------
    return render(request, "student/review/3_ma.html", {
        "student_name": student_name,
        "username": username,
        "team_name": team_title,
        "uploaded_files": uploaded_files,
        "remarks_by_type": remarks_by_type,
        "highlighted_pdfs": highlighted_pdfs,
    })

import json
from django.http import JsonResponse
from .models import ZerothReviewRemark, AllocationResult

def save_zeroth_remark(request):
    """
    Save remarks for a specific file type (abstract/pdf/ppt).
    Each remark is tied to a specific file type and will only show 
    when that file type is viewed.
    """
    print("🔥 save_zeroth_remark CALLED")

    if request.method != "POST":
        return JsonResponse({"status": "fail", "message": "Invalid request"})

    mentor_name = request.session.get("mentor_name")
    print("Mentor:", mentor_name)

    allocation = AllocationResult.objects.filter(
        mentor_name=mentor_name
    ).first()

    if not allocation:
        print("❌ No allocation found")
        return JsonResponse({"status": "fail", "message": "No team allocated"})

    team_name = allocation.team_name
    print("Team:", team_name)

    try:
        data = json.loads(request.body)
        remarks = data.get("remarks", [])
        deleted = data.get("deleted", [])
        
        file_type = data.get("file_type", "abstract")
        
        if file_type not in ["abstract", "pdf", "ppt"]:
            print(f"⚠️ Invalid file_type '{file_type}', using 'abstract'")
            file_type = "abstract"
            
        print("File type:", file_type)
        print("Remarks count:", len(remarks))
        print("Deleted count:", len(deleted))
        
    except Exception as e:
        print("❌ JSON error:", e)
        return JsonResponse({"status": "fail", "message": "Invalid JSON"})

    inserted = 0
    updated = 0
    deleted_count = 0

    # Handle deletions
    if deleted:
        for heading in deleted:
            heading = heading.strip()
            if not heading:
                continue
                
            print(f"🗑️ Deleting: '{heading}' for {file_type}")
            
            deleted_count += ZerothReviewRemark.objects.filter(
                team_name=team_name,
                mentor_name=mentor_name,
                heading=heading,
                file_type=file_type
            ).delete()[0]

    # Handle upserts
    for r in remarks:
        heading = (r.get("heading") or "").strip()
        remark = (r.get("remark") or "").strip()
        color = r.get("color") or "#ffe066"
        
        # 🔥 CRITICAL FIX: Extract coordinates from payload
        coordinates = r.get("coordinates") or {}
        print(f"🔥 Raw coordinates from frontend: {coordinates}")

        if not heading or not remark:
            continue

        # 🔥 CRITICAL FIX: Convert coordinates dict to JSON string for TextField
        if isinstance(coordinates, dict):
            coords_json = json.dumps(coordinates)
        elif isinstance(coordinates, str):
            coords_json = coordinates
        else:
            coords_json = "{}"
        
        print(f"💾 Saving: '{heading}' for {file_type}")
        print(f"🔥 Coordinates JSON (first 100 chars): {coords_json[:100]}")

        obj, created = ZerothReviewRemark.objects.update_or_create(
            team_name=team_name,
            mentor_name=mentor_name,
            heading=heading,
            file_type=file_type,
            defaults={
                "remark": remark,
                "color": color,
                "coordinates": coords_json,  # 🔥 FIXED: Now saving coordinates!
            }
        )

        print(f"✅ Saved record ID={obj.id}, coordinates in DB={obj.coordinates[:100] if obj.coordinates else 'EMPTY'}")

        if created:
            inserted += 1
        else:
            updated += 1

    print(f"✅ Done for {file_type}: {inserted} new, {updated} updated, {deleted_count} deleted")
    
    return JsonResponse({
        "status": "success",
        "file_type": file_type,
        "inserted": inserted,
        "updated": updated,
        "deleted": deleted_count
    })


from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json

from allocation.models import ZerothReviewRemark, AllocationResult

from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

@csrf_exempt
def save_zeroth_evaluation(request):
    if request.method != "POST":
        return JsonResponse({"status": "fail", "message": "Invalid request"})

    try:
        data = json.loads(request.body)
        team_name = data.get("team_name")
        evaluations = data.get("evaluations", {})
        
        print(f"Received team_name: {team_name}")
        print(f"Received evaluations: {evaluations}")
        
        doc = generate_zero_review_docx_from_template(team_name, evaluations)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        safe_name = team_name.replace(" ", "_")
        filename = f"{safe_name}_ZerothReview.docx"
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return JsonResponse({"status": "fail", "message": str(e)})


def get_cell_text(cell):
    return cell.text.strip()


def set_cell_text(cell, text, bold=False, size=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def generate_zero_review_docx_from_template(team_name, evaluations):
    template_path = os.path.join(settings.BASE_DIR, 'static', 'zeroth_review_mark.docx')
    if not os.path.exists(template_path):
        template_path = os.path.join(settings.BASE_DIR, 'allocation', 'static', 'zeroth_review_mark.docx')
    if not os.path.exists(template_path):
        template_path = os.path.join(settings.BASE_DIR, 'mentor', 'static', 'zeroth_review_mark.docx')
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at: {template_path}")
    
    doc = Document(template_path)
    
    # Extract members from evaluations and sort to maintain order
    members = []
    for key in evaluations.keys():
        members.append(key.replace('team_member-', ''))
    members.sort()
    
    print(f"Members: {members}")
    print(f"Evaluations: {evaluations}")
    
    criteria_order = ["concept", "literature", "impact", "planning", "methodology", "presentation"]
    criteria_labels = [
        "Project Concept & Topic Selection(20)",
        "Literature Review & Background Research (10)",
        "Relevance & Impact of Project Outcome (20)",
        "Project Planning & Timeline (20)",
        "Methodology & Approach (20)",
        "Presentation & Clarity (10)"
    ]
    
    safe_name = team_name.replace(" ", "_")
    title_text = f"{safe_name}_ZerothReview"
    today = datetime.now().strftime("%d/%m/%Y")
    
    # === STEP 1: Fill paragraphs ===
    title_inserted = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Insert title after "ZEROTH REVIEW" paragraph
        if not title_inserted and "ZEROTH REVIEW" in text.upper():
            new_p = doc.add_paragraph(title_text)
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_p.runs[0].bold = True
            new_p.runs[0].font.size = Pt(14)
            para._element.addnext(new_p._element)
            title_inserted = True
        
        # Fill Date
        if text.startswith("Date:"):
            para.clear()
            para.add_run("Date: ").bold = True
            para.add_run(today)
        
        # Fill Project Title
        if text.startswith("Project Title:"):
            para.clear()
            para.add_run("Project Title: ").bold = True
            para.add_run(team_name)
    
    # === STEP 2: Fill Table 0 - TEAM MEMBERS ===
    table0 = doc.tables[0]
    for m_idx, member in enumerate(members):
        row_idx = m_idx + 2  # Skip merged header row 0 and header row 1
        if row_idx < len(table0.rows):
            row = table0.rows[row_idx]
            set_cell_text(row.cells[0], str(m_idx + 1))  # S.NO.
            if len(row.cells) > 3:
                set_cell_text(row.cells[3], member)  # STUDENT NAME
    
    # === STEP 3: Fill Table 1 - MARKS TABLE ===
    table1 = doc.tables[1]
    
    # Find data start row
    data_start_row = -1
    for r_idx, row in enumerate(table1.rows):
        first_text = get_cell_text(row.cells[0]).strip()
        if "Project Concept" in first_text:
            data_start_row = r_idx
            break
    
    if data_start_row == -1:
        data_start_row = 5
    
    print(f"Data starts at row {data_start_row}")
    
    criteria_idx = 0
    for r_idx in range(data_start_row, len(table1.rows)):
        row = table1.rows[r_idx]
        first_cell_text = get_cell_text(row.cells[0]).strip()
        
        print(f"Row {r_idx}: '{first_cell_text}' | cells: {len(row.cells)}")
        
        if not first_cell_text:
            continue
        
        # TOTAL row
        if "TOTAL" in first_cell_text.upper():
            print(f"Found TOTAL row")
            for m_idx, member in enumerate(members):
                member_key = f"team_member-{member}"
                marks_list = evaluations.get(member_key, [])
                total = sum(int(m) if str(m).isdigit() else 0 for m in marks_list)
                print(f"  Member {member} total: {total}")
                
                # Try columns 4,5,6,7 first
                col_idx = 4 + m_idx
                if col_idx < len(row.cells):
                    set_cell_text(row.cells[col_idx], total, bold=True)
                else:
                    # Fallback: find first empty cell after column 1
                    for c in range(2, len(row.cells)):
                        if get_cell_text(row.cells[c]) == "":
                            set_cell_text(row.cells[c], total, bold=True)
                            break
            break
        
        # Match criteria
        matched_idx = -1
        for c_idx, label in enumerate(criteria_labels):
            if label.lower() in first_cell_text.lower():
                matched_idx = c_idx
                break
        
        if matched_idx == -1:
            partial_map = {
                "concept": ["concept", "topic selection", "aim"],
                "literature": ["literature", "background", "knowledge about existing"],
                "impact": ["relevance", "impact", "problem identification"],
                "planning": ["planning", "timeline", "technical design"],
                "methodology": ["methodology", "approach", "organization of the presentation"],
                "presentation": ["presentation", "clarity", "estimate of the proposed"]
            }
            for c_idx, criteria_key in enumerate(criteria_order):
                for keyword in partial_map.get(criteria_key, []):
                    if keyword in first_cell_text.lower():
                        matched_idx = c_idx
                        break
                if matched_idx != -1:
                    break
        
        if matched_idx != -1:
            print(f"  Matched criteria idx {matched_idx}")
            
            for m_idx, member in enumerate(members):
                member_key = f"team_member-{member}"
                marks_list = evaluations.get(member_key, [])
                mark = marks_list[matched_idx] if matched_idx < len(marks_list) else "0"
                print(f"    Member {member} ({member_key}): mark={mark}")
                
                col_idx = 4 + m_idx
                if col_idx < len(row.cells):
                    set_cell_text(row.cells[col_idx], mark)
            
            # Fill max marks
            max_marks = [20, 10, 20, 20, 20, 10][matched_idx]
            if len(row.cells) > 2:
                set_cell_text(row.cells[2], max_marks)
            
            criteria_idx += 1
    
    return doc

def save_first_remark(request):
    """
    Save remarks for PPT in First Review.
    """
    print("🔥 save_first_remark CALLED")

    if request.method != "POST":
        return JsonResponse({"status": "fail", "message": "Invalid request"})

    mentor_name = request.session.get("mentor_name")
    print("Mentor:", mentor_name)

    allocation = AllocationResult.objects.filter(
        mentor_name=mentor_name
    ).first()

    if not allocation:
        print("❌ No allocation found")
        return JsonResponse({"status": "fail", "message": "No team allocated"})

    team_name = allocation.team_name
    print("Team:", team_name)

    try:
        data = json.loads(request.body)
        remarks = data.get("remarks", [])
        deleted = data.get("deleted", [])
        
        # FIXED: Only PPT for first review
        file_type = "ppt"
            
        print("File type:", file_type)
        print("Remarks count:", len(remarks))
        print("Deleted count:", len(deleted))
        
    except Exception as e:
        print("❌ JSON error:", e)
        return JsonResponse({"status": "fail", "message": "Invalid JSON"})

    inserted = 0
    updated = 0
    deleted_count = 0

    # Handle deletions
    if deleted:
        for heading in deleted:
            heading = heading.strip()
            if not heading:
                continue
                
            print(f"🗑️ Deleting: '{heading}' for {file_type}")
            
            deleted_count += FirstReviewRemark.objects.filter(
                team_name=team_name,
                mentor_name=mentor_name,
                heading=heading,
                file_type=file_type
            ).delete()[0]

    # Handle upserts
    for r in remarks:
        heading = (r.get("heading") or "").strip()
        remark = (r.get("remark") or "").strip()
        color = r.get("color") or "#ffe066"
        slide_number = r.get("slideNumber") or None

        if not heading or not remark:
            continue

        print(f"💾 Saving: '{heading}' for {file_type}")

        obj, created = FirstReviewRemark.objects.update_or_create(
            team_name=team_name,
            mentor_name=mentor_name,
            heading=heading,
            file_type=file_type,
            defaults={
                "remark": remark,
                "color": color,
            }
        )

        if created:
            inserted += 1
        else:
            updated += 1

    print(f"✅ Done for First Review {file_type}: {inserted} new, {updated} updated, {deleted_count} deleted")
    
    return JsonResponse({
        "status": "success",
        "file_type": file_type,
        "inserted": inserted,
        "updated": updated,
        "deleted": deleted_count
    })


def save_second_remark(request):
    """
    Save remarks for PPT in Second Review.
    """
    print("🔥 save_second_remark CALLED")

    if request.method != "POST":
        return JsonResponse({"status": "fail", "message": "Invalid request"})

    mentor_name = request.session.get("mentor_name")
    print("Mentor:", mentor_name)

    allocation = AllocationResult.objects.filter(
        mentor_name=mentor_name
    ).first()

    if not allocation:
        print("❌ No allocation found")
        return JsonResponse({"status": "fail", "message": "No team allocated"})

    team_name = allocation.team_name
    print("Team:", team_name)

    try:
        data = json.loads(request.body)
        remarks = data.get("remarks", [])
        deleted = data.get("deleted", [])
        
        # FIXED: Only PPT for second review
        file_type = "ppt"
            
        print("File type:", file_type)
        print("Remarks count:", len(remarks))
        print("Deleted count:", len(deleted))
        
    except Exception as e:
        print("❌ JSON error:", e)
        return JsonResponse({"status": "fail", "message": "Invalid JSON"})

    inserted = 0
    updated = 0
    deleted_count = 0

    # Handle deletions
    if deleted:
        for heading in deleted:
            heading = heading.strip()
            if not heading:
                continue
                
            print(f"🗑️ Deleting: '{heading}' for {file_type}")
            
            deleted_count += SecondReviewRemark.objects.filter(
                team_name=team_name,
                mentor_name=mentor_name,
                heading=heading,
                file_type=file_type
            ).delete()[0]

    # Handle upserts
    for r in remarks:
        heading = (r.get("heading") or "").strip()
        remark = (r.get("remark") or "").strip()
        color = r.get("color") or "#ffe066"
        slide_number = r.get("slideNumber") or None

        if not heading or not remark:
            continue

        print(f"💾 Saving: '{heading}' for {file_type}")

        obj, created = SecondReviewRemark.objects.update_or_create(
            team_name=team_name,
            mentor_name=mentor_name,
            heading=heading,
            file_type=file_type,
            defaults={
                "remark": remark,
                "color": color,
            }
        )

        if created:
            inserted += 1
        else:
            updated += 1

    print(f"✅ Done for Second Review {file_type}: {inserted} new, {updated} updated, {deleted_count} deleted")
    
    return JsonResponse({
        "status": "success",
        "file_type": file_type,
        "inserted": inserted,
        "updated": updated,
        "deleted": deleted_count
    })


def save_third_remark(request):
    """
    Save remarks for PPT in Third Review.
    """
    print("🔥 save_third_remark CALLED")

    if request.method != "POST":
        return JsonResponse({"status": "fail", "message": "Invalid request"})

    mentor_name = request.session.get("mentor_name")
    print("Mentor:", mentor_name)

    allocation = AllocationResult.objects.filter(
        mentor_name=mentor_name
    ).first()

    if not allocation:
        print("❌ No allocation found")
        return JsonResponse({"status": "fail", "message": "No team allocated"})

    team_name = allocation.team_name
    print("Team:", team_name)

    try:
        data = json.loads(request.body)
        remarks = data.get("remarks", [])
        deleted = data.get("deleted", [])
        
        # FIXED: Only PPT for third review
        file_type = "ppt"
            
        print("File type:", file_type)
        print("Remarks count:", len(remarks))
        print("Deleted count:", len(deleted))
        
    except Exception as e:
        print("❌ JSON error:", e)
        return JsonResponse({"status": "fail", "message": "Invalid JSON"})

    inserted = 0
    updated = 0
    deleted_count = 0

    # Handle deletions
    if deleted:
        for heading in deleted:
            heading = heading.strip()
            if not heading:
                continue
                
            print(f"🗑️ Deleting: '{heading}' for {file_type}")
            
            deleted_count += ThirdReviewRemark.objects.filter(
                team_name=team_name,
                mentor_name=mentor_name,
                heading=heading,
                file_type=file_type
            ).delete()[0]

    # Handle upserts
    for r in remarks:
        heading = (r.get("heading") or "").strip()
        remark = (r.get("remark") or "").strip()
        color = r.get("color") or "#ffe066"
        slide_number = r.get("slideNumber") or None

        if not heading or not remark:
            continue

        print(f"💾 Saving: '{heading}' for {file_type}")

        obj, created = ThirdReviewRemark.objects.update_or_create(
            team_name=team_name,
            mentor_name=mentor_name,
            heading=heading,
            file_type=file_type,
            defaults={
                "remark": remark,
                "color": color,
            }
        )

        if created:
            inserted += 1
        else:
            updated += 1

    print(f"✅ Done for Third Review {file_type}: {inserted} new, {updated} updated, {deleted_count} deleted")
    
    return JsonResponse({
        "status": "success",
        "file_type": file_type,
        "inserted": inserted,
        "updated": updated,
        "deleted_count": deleted_count
    })

def clean_text(text):
    return re.sub(r'\(.*?\)', '', text).strip().lower()


@csrf_exempt
def save_evaluation(request):
    """
    📝 Save Zeroth Review Evaluation Marks into DOCX
    """

    if request.method != "POST":
        return JsonResponse(
            {"status": "fail", "message": "Invalid request method"},
            status=400
        )

    try:
        data = json.loads(request.body)
        team_name = data.get("team_name")
        evaluations = data.get("evaluations")  # dict

        if not team_name or not evaluations:
            return JsonResponse(
                {"status": "fail", "message": "Missing team name or evaluations"},
                status=400
            )

        # -------------------------------------------------
        # Safe team name (filesystem)
        # -------------------------------------------------
        team_name_fs = team_name.replace(" ", "_")

        # -------------------------------------------------
        # Paths
        # -------------------------------------------------
        template_path = os.path.join(
            settings.BASE_DIR,
            "allocation",
            "static",
            "zeroth_review_mark.docx"
        )

        output_dir = os.path.join(settings.BASE_DIR, "generated_docs")
        os.makedirs(output_dir, exist_ok=True)

        output_path = os.path.join(
            output_dir,
            f"{team_name_fs}_ZerothReview.docx"
        )

        print("[DEBUG] Output DOCX:", output_path)

        # -------------------------------------------------
        # Load existing doc OR template
        # -------------------------------------------------
        if os.path.exists(output_path):
            doc = Document(output_path)
        else:
            if not os.path.exists(template_path):
                return JsonResponse(
                    {
                        "status": "error",
                        "message": f"Template not found: {template_path}"
                    },
                    status=500
                )
            doc = Document(template_path)

        # -------------------------------------------------
        # Insert project title
        # -------------------------------------------------
        for para in doc.paragraphs:
            if "project title" in para.text.lower():
                para.text = f"Project Title: {team_name}"
                break

        # -------------------------------------------------
        # Locate Team Members table
        # -------------------------------------------------
        members_table = None
        for t in doc.tables:
            if "team members" in clean_text(t.cell(0, 0).text):
                members_table = t
                break

        if not members_table:
            return JsonResponse(
                {"status": "error", "message": "Team Members table not found"},
                status=500
            )

        # -------------------------------------------------
        # Existing members
        # -------------------------------------------------
        existing_names = []
        for r in members_table.rows[1:]:
            if len(r.cells) > 3 and r.cells[3].text.strip():
                existing_names.append(r.cells[3].text.strip())

        current_index = len(existing_names) + 1  # ✅ start from next S.No

        for member_key in evaluations.keys():
            clean_name = member_key.replace("team_member-", "").strip()
            if clean_name in existing_names:
                continue

            empty_row = next(
                (r for r in members_table.rows[1:] if not r.cells[3].text.strip()),
                None
            )

            if not empty_row:
                empty_row = members_table.add_row()
                for c in empty_row.cells:
                    c.text = ""

            empty_row.cells[0].text = str(current_index)
            empty_row.cells[1].text = "-"
            empty_row.cells[2].text = "-"
            empty_row.cells[3].text = clean_name

            existing_names.append(clean_name)
            current_index += 1

        # -------------------------------------------------
        # Locate Marks table
        # -------------------------------------------------
        marks_table = None
        for t in doc.tables:
            for row in t.rows:
                if any(
                    k in clean_text(row.cells[0].text)
                    for k in [
                        "project concept",
                        "literature review",
                        "relevance",
                        "project planning",
                        "methodology",
                        "presentation"
                    ]
                ):
                    marks_table = t
                    break
            if marks_table:
                break

        if not marks_table:
            return JsonResponse(
                {"status": "error", "message": "Marks table not found"},
                status=500
            )

        # -------------------------------------------------
        # Criteria map + total row
        # -------------------------------------------------
        criteria_map = {}
        total_row = None

        for i, row in enumerate(marks_table.rows):
            t0 = clean_text(row.cells[0].text)

            if "total" in t0:
                total_row = i

            for key in [
                "project concept",
                "literature review",
                "relevance",
                "project planning",
                "methodology",
                "presentation"
            ]:
                if key in t0:
                    criteria_map[key] = i

        if total_row is None:
            return JsonResponse(
                {"status": "error", "message": "Total row not found"},
                status=500
            )

        # -------------------------------------------------
        # Member → S.No map
        # -------------------------------------------------
        member_to_sno = {}
        for r in members_table.rows[1:]:
            if len(r.cells) > 3 and r.cells[3].text.strip():
                name = r.cells[3].text.strip().lower()
                sno = r.cells[0].text.strip()
                member_to_sno[name] = sno

        print("[DEBUG] Member → S.No:", member_to_sno)

        # -------------------------------------------------
        # Insert marks
        # -------------------------------------------------
        for member_key, marks_list in evaluations.items():
            name = member_key.replace("team_member-", "").strip().lower()
            sno = member_to_sno.get(name)

            if not sno:
                continue

            col = 3 + int(sno)  # ✅ correct column mapping

            if col >= len(marks_table.rows[0].cells):
                continue

            marks_dict = {}
            for item in marks_list:
                if "-" in item:
                    crit, val = item.rsplit("-", 1)
                    try:
                        marks_dict[crit.strip()] = int(val.lstrip("0") or "0")
                    except:
                        pass

            total = 0

            for crit, mark in marks_dict.items():
                ckey = clean_text(crit)
                row_index = next(
                    (
                        criteria_map[k]
                        for k in criteria_map
                        if k in ckey or ckey in k
                    ),
                    None
                )

                if row_index is not None:
                    marks_table.rows[row_index].cells[col].text = str(mark)
                    total += mark

            marks_table.rows[total_row].cells[col].text = str(total)

        # -------------------------------------------------
        # Safe save
        # -------------------------------------------------
        try:
            doc.save(output_path)
        except PermissionError:
            ts = time.strftime("%Y%m%d_%H%M%S")
            alt_path = os.path.join(
                output_dir,
                f"{team_name_fs}_ZerothReview_{ts}.docx"
            )
            doc.save(alt_path)
            output_path = alt_path

        return JsonResponse(
            {
                "status": "success",
                "message": "Marks inserted successfully",
                "file_path": output_path
            }
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse(
            {"status": "error", "message": str(e)},
            status=500
        )


from django.http import FileResponse, JsonResponse # type: ignore
import os
from django.conf import settings


try:
    import pdfkit
    from docx2pdf import convert
except ImportError:
    pdfkit = None
    convert = None

from django.shortcuts import redirect
from django.http import JsonResponse

def download_docx(request, team_name):
    """
    📥 Download Zeroth Review DOCX (CLOUDINARY ONLY)
    """

    if not team_name:
        return JsonResponse(
            {"status": "fail", "message": "Invalid team name"},
            status=400
        )

    # -------------------------------------------------
    # Get allocation by team name
    # -------------------------------------------------
    allocation = AllocationResult.objects.filter(
        team_name=team_name
    ).first()

    if not allocation:
        return JsonResponse(
            {"status": "fail", "message": "Team not found"},
            status=404
        )

    # -------------------------------------------------
    # Cloudinary DOCX URL
    # -------------------------------------------------
    docx_url = allocation.zeroth_review_docx_url

    if not docx_url:
        return JsonResponse(
            {
                "status": "fail",
                "message": "DOCX not uploaded to Cloudinary"
            },
            status=404
        )

    print(f"[DEBUG] Redirecting to Cloudinary DOCX: {docx_url}")

    # -------------------------------------------------
    # Redirect to Cloudinary (download handled by Cloudinary)
    # -------------------------------------------------
    return redirect(docx_url)


import os
import pdfkit
import os
try:
    import pdfkit
    from docx2pdf import convert
except ImportError:
    pdfkit = None
    convert = None
from django.http import FileResponse, JsonResponse
from django.conf import settings
from docx import Document
from tempfile import NamedTemporaryFile

def download_pdf(request, team_name):
    """
    📥 Download Zeroth Review PDF (CLOUDINARY ONLY)
    """

    if not team_name:
        return JsonResponse(
            {"status": "fail", "message": "Invalid team name"},
            status=400
        )

    # -------------------------------------------------
    # Fetch allocation
    # -------------------------------------------------
    allocation = AllocationResult.objects.filter(
        team_name=team_name
    ).first()

    if not allocation:
        return JsonResponse(
            {"status": "fail", "message": "Team not found"},
            status=404
        )

    # -------------------------------------------------
    # Cloudinary PDF URL
    # -------------------------------------------------
    pdf_url = allocation.zeroth_review_pdf_url

    if not pdf_url:
        return JsonResponse(
            {
                "status": "fail",
                "message": "PDF not uploaded to Cloudinary"
            },
            status=404
        )

    print(f"[DEBUG] Redirecting to Cloudinary PDF: {pdf_url}")

    # -------------------------------------------------
    # Redirect (Cloudinary handles download)
    # -------------------------------------------------
    return redirect(pdf_url)

import os, json, time
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from docx import Document

def clean_text(text):
    return text.strip().lower()


import json
import os
import time
from django.http import JsonResponse
from django.conf import settings
from docx import Document
from django.views.decorators.csrf import csrf_exempt


from django.http import JsonResponse, HttpResponse
from docx import Document
from django.conf import settings
import json, os, io, time, traceback

@csrf_exempt
def save_evaluation_review1(request):
    """
    📝 Save FIRST REVIEW Evaluation Marks
    📥 DIRECT DOCX DOWNLOAD (NO LOCAL SAVE)
    """

    if request.method != "POST":
        return JsonResponse(
            {"status": "fail", "message": "Invalid request method"},
            status=400
        )

    try:
        data = json.loads(request.body)
        team_name = data.get("team_name")
        evaluations = data.get("evaluations")  # {"team_member-X": [marks]}

        if not team_name or not evaluations:
            return JsonResponse(
                {"status": "fail", "message": "Missing team name or evaluations"},
                status=400
            )

        team_name_fs = team_name.replace(" ", "_")

        # -------------------------------------------------
        # Load DOCX template
        # -------------------------------------------------
        template_path = os.path.join(
            settings.BASE_DIR,
            "allocation",
            "static",
            "first_review_mark.docx"
        )

        if not os.path.exists(template_path):
            return JsonResponse(
                {"status": "error", "message": "DOCX template not found"},
                status=500
            )

        doc = Document(template_path)

        # -------------------------------------------------
        # Update title
        # -------------------------------------------------
        for para in doc.paragraphs:
            if "review 1" in para.text.lower():
                para.text = f"Review 1 Evaluation - {team_name}"
                break

        # -------------------------------------------------
        # TEAM MEMBERS TABLE (Table 0)
        # -------------------------------------------------
        members_table = doc.tables[0]
        start_row = 2
        existing_names = []

        for r in members_table.rows[start_row:]:
            if len(r.cells) >= 4 and r.cells[3].text.strip():
                existing_names.append(r.cells[3].text.strip())

        current_index = len(existing_names) + 1

        for member_key in evaluations.keys():
            name = member_key.replace("team_member-", "").strip()

            if name in existing_names:
                continue

            if start_row + current_index - 1 >= len(members_table.rows):
                members_table.add_row()

            row = members_table.rows[start_row + current_index - 1]
            row.cells[0].text = str(current_index)
            row.cells[3].text = name

            existing_names.append(name)
            current_index += 1

        # -------------------------------------------------
        # Map MEMBER → S.NO
        # -------------------------------------------------
        member_to_sno = {}
        for r in members_table.rows[start_row:]:
            if len(r.cells) >= 4 and r.cells[3].text.strip():
                member_to_sno[
                    r.cells[3].text.strip().lower()
                ] = r.cells[0].text.strip()

        # -------------------------------------------------
        # MARKS TABLE (Table 1)
        # -------------------------------------------------
        marks_table = doc.tables[1]

        # Detect TOTAL row
        total_row = None
        for i, row in enumerate(marks_table.rows):
            if "total" in row.cells[0].text.lower():
                total_row = i
                break

        if total_row is None:
            total_row = len(marks_table.rows) - 1

        # Detect S.NO → column mapping
        sno_col_map = {}
        sno_row_idx = None

        for i, row in enumerate(marks_table.rows):
            for idx, cell in enumerate(row.cells):
                if cell.text.strip().isdigit():
                    sno_col_map[cell.text.strip()] = idx
                    sno_row_idx = i
            if sno_col_map:
                break

        print("[DEBUG] S.NO → Column:", sno_col_map)

        # -------------------------------------------------
        # Insert marks
        # -------------------------------------------------
        for member_key, marks_list in evaluations.items():
            name = member_key.replace("team_member-", "").strip().lower()
            sno = member_to_sno.get(name)

            if not sno:
                continue

            col_idx = sno_col_map.get(sno)
            if col_idx is None:
                continue

            total = 0
            row_idx = sno_row_idx + 1

            for mark in marks_list:
                if row_idx >= total_row:
                    break
                try:
                    marks_table.rows[row_idx].cells[col_idx].text = str(mark)
                    total += int(mark)
                except Exception:
                    pass
                row_idx += 1

            marks_table.rows[total_row].cells[col_idx].text = str(total)

        # -------------------------------------------------
        # STREAM DOCX (NO LOCAL SAVE)
        # -------------------------------------------------
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = (
            f'attachment; filename="{team_name_fs}_Review1.docx"'
        )

        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse(
            {"status": "error", "message": str(e)},
            status=500
        )

def save_evaluation_review2(request):
    """
    📝 Save Second REVIEW Evaluation Marks into DOCX
    """

    if request.method != "POST":
        return JsonResponse(
            {"status": "fail", "message": "Invalid request method"},
            status=400
        )

    try:
        data = json.loads(request.body)
        team_name = data.get("team_name")
        evaluations = data.get("evaluations")  # {"member_name": [marks list]}

        if not team_name or not evaluations:
            return JsonResponse(
                {"status": "fail", "message": "Missing team name or evaluations"},
                status=400
            )

        # -------------------------------------------------
        # Safe team name (filesystem)
        # -------------------------------------------------
        team_name_fs = team_name.replace(" ", "_")

        # -------------------------------------------------
        # Paths
        # -------------------------------------------------
        template_path = os.path.join(
            settings.BASE_DIR,
            "allocation",
            "static",
            "second_review_mark.docx"
        )

        output_dir = os.path.join(settings.BASE_DIR, "generated_docs")
        os.makedirs(output_dir, exist_ok=True)

        output_path = os.path.join(
            output_dir,
            f"{team_name_fs}_Review2.docx"
        )

        print("[DEBUG] Review2 output:", output_path)

        # -------------------------------------------------
        # Load existing doc OR template
        # -------------------------------------------------
        if os.path.exists(output_path):
            doc = Document(output_path)
        else:
            if not os.path.exists(template_path):
                return JsonResponse(
                    {
                        "status": "error",
                        "message": f"Template not found: {template_path}"
                    },
                    status=500
                )
            doc = Document(template_path)

        # -------------------------------------------------
        # Update title
        # -------------------------------------------------
        for para in doc.paragraphs:
            if "review 2" in para.text.lower():
                para.text = f"Review 2 Evaluation - {team_name}"
                break

        # -------------------------------------------------
        # TEAM MEMBERS table (assumed first table)
        # -------------------------------------------------
        members_table = doc.tables[0]

        start_row = 2  # after headers
        existing_names = []

        for r in members_table.rows[start_row:]:
            if len(r.cells) >= 4 and r.cells[3].text.strip():
                existing_names.append(r.cells[3].text.strip())

        current_index = len(existing_names) + 1

        for member_key in evaluations.keys():
            clean_name = member_key.replace("team_member-", "").strip()

            if clean_name in existing_names:
                continue

            row_index = start_row + (current_index - 1)
            if row_index >= len(members_table.rows):
                members_table.add_row()

            members_table.rows[row_index].cells[0].text = str(current_index)
            members_table.rows[row_index].cells[3].text = clean_name

            existing_names.append(clean_name)
            current_index += 1

        # -------------------------------------------------
        # Map member → S.NO
        # -------------------------------------------------
        member_to_sno = {}
        for r in members_table.rows[start_row:]:
            if len(r.cells) >= 4 and r.cells[3].text.strip():
                name = r.cells[3].text.strip().lower()
                sno = r.cells[0].text.strip()
                member_to_sno[name] = sno

        print("[DEBUG] Member → S.NO:", member_to_sno)

        # -------------------------------------------------
        # MARKS table (assumed second table)
        # -------------------------------------------------
        marks_table = doc.tables[1]

        # -------------------------------------------------
        # Detect TOTAL row
        # -------------------------------------------------
        total_row = None
        for i, row in enumerate(marks_table.rows):
            if "total" in row.cells[0].text.lower():
                total_row = i
                break

        if total_row is None:
            total_row = len(marks_table.rows) - 1

        # -------------------------------------------------
        # Detect S.NO → column mapping
        # -------------------------------------------------
        sno_col_map = {}
        sno_row_idx = None

        for i, row in enumerate(marks_table.rows):
            for idx, cell in enumerate(row.cells):
                if cell.text.strip().isdigit():
                    sno_col_map[cell.text.strip()] = idx
                    sno_row_idx = i
            if sno_col_map:
                break

        print("[DEBUG] S.NO → Column:", sno_col_map)

        # -------------------------------------------------
        # Insert marks
        # -------------------------------------------------
        for member_key, marks_list in evaluations.items():
            clean_name = member_key.replace("team_member-", "").strip().lower()
            sno = member_to_sno.get(clean_name)

            if not sno:
                continue

            col_idx = sno_col_map.get(sno)
            if col_idx is None:
                continue

            total = 0
            row_idx = sno_row_idx + 1

            for mark in marks_list:
                if row_idx >= total_row:
                    break
                try:
                    marks_table.rows[row_idx].cells[col_idx].text = str(mark)
                    total += int(mark)
                except:
                    pass
                row_idx += 1

            marks_table.rows[total_row].cells[col_idx].text = str(total)

        # -------------------------------------------------
        # Safe save
        # -------------------------------------------------
        try:
            doc.save(output_path)
        except PermissionError:
            ts = time.strftime("%Y%m%d_%H%M%S")
            alt = os.path.join(
                output_dir,
                f"{team_name_fs}_Review2_{ts}.docx"
            )
            doc.save(alt)
            output_path = alt

        return JsonResponse(
            {
                "status": "success",
                "message": "Review 2 marks saved successfully",
                "file_path": output_path
            }
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse(
            {"status": "error", "message": str(e)},
            status=500
        )

def save_evaluation_review3(request):
    """
    📝 Save THIRD REVIEW Evaluation Marks into DOCX
    """
    print("hello")
    if request.method != "POST":
        return JsonResponse(
            {"status": "fail", "message": "Invalid request method"},
            status=400
        )

    try:
        data = json.loads(request.body)
        team_name = data.get("team_name")
        evaluations = data.get("evaluations")  # {"member_name": [marks list]}

        if not team_name or not evaluations:
            return JsonResponse(
                {"status": "fail", "message": "Missing team name or evaluations"},
                status=400
            )

        # -------------------------------------------------
        # Safe team name (filesystem)
        # -------------------------------------------------
        team_name_fs = team_name.replace(" ", "_")

        # -------------------------------------------------
        # Paths
        # -------------------------------------------------
        template_path = os.path.join(
            settings.BASE_DIR,
            "allocation",
            "static",
            "third_review_mark.docx"
        )

        output_dir = os.path.join(settings.BASE_DIR, "generated_docs")
        os.makedirs(output_dir, exist_ok=True)

        output_path = os.path.join(
            output_dir,
            f"{team_name_fs}_Review3.docx"
        )

        print("[DEBUG] Review3 output:", output_path)

        # -------------------------------------------------
        # Load existing doc OR template
        # -------------------------------------------------
        if os.path.exists(output_path):
            doc = Document(output_path)
        else:
            if not os.path.exists(template_path):
                return JsonResponse(
                    {
                        "status": "error",
                        "message": f"Template not found: {template_path}"
                    },
                    status=500
                )
            doc = Document(template_path)

        # -------------------------------------------------
        # Update title
        # -------------------------------------------------
        for para in doc.paragraphs:
            if "review 3" in para.text.lower():
                para.text = f"Review 3 Evaluation - {team_name}"
                break

        # -------------------------------------------------
        # TEAM MEMBERS table (assumed first table)
        # -------------------------------------------------
        members_table = doc.tables[0]

        start_row = 2  # after headers
        existing_names = []

        for r in members_table.rows[start_row:]:
            if len(r.cells) >= 4 and r.cells[3].text.strip():
                existing_names.append(r.cells[3].text.strip())

        current_index = len(existing_names) + 1

        for member_key in evaluations.keys():
            clean_name = member_key.replace("team_member-", "").strip()

            if clean_name in existing_names:
                continue

            row_index = start_row + (current_index - 1)
            if row_index >= len(members_table.rows):
                members_table.add_row()

            members_table.rows[row_index].cells[0].text = str(current_index)
            members_table.rows[row_index].cells[3].text = clean_name

            existing_names.append(clean_name)
            current_index += 1

        # -------------------------------------------------
        # Map member → S.NO
        # -------------------------------------------------
        member_to_sno = {}
        for r in members_table.rows[start_row:]:
            if len(r.cells) >= 4 and r.cells[3].text.strip():
                name = r.cells[3].text.strip().lower()
                sno = r.cells[0].text.strip()
                member_to_sno[name] = sno

        print("[DEBUG] Member → S.NO:", member_to_sno)

        # -------------------------------------------------
        # MARKS table (assumed second table)
        # -------------------------------------------------
        marks_table = doc.tables[1]

        # -------------------------------------------------
        # Detect TOTAL row
        # -------------------------------------------------
        total_row = None
        for i, row in enumerate(marks_table.rows):
            if "total" in row.cells[0].text.lower():
                total_row = i
                break

        if total_row is None:
            total_row = len(marks_table.rows) - 1

        # -------------------------------------------------
        # Detect S.NO → column mapping
        # -------------------------------------------------
        sno_col_map = {}
        sno_row_idx = None

        for i, row in enumerate(marks_table.rows):
            for idx, cell in enumerate(row.cells):
                if cell.text.strip().isdigit():
                    sno_col_map[cell.text.strip()] = idx
                    sno_row_idx = i
            if sno_col_map:
                break

        print("[DEBUG] S.NO → Column:", sno_col_map)

        # -------------------------------------------------
        # Insert marks
        # -------------------------------------------------
        for member_key, marks_list in evaluations.items():
            clean_name = member_key.replace("team_member-", "").strip().lower()
            sno = member_to_sno.get(clean_name)

            if not sno:
                continue

            col_idx = sno_col_map.get(sno)
            if col_idx is None:
                continue

            total = 0
            row_idx = sno_row_idx + 1

            for mark in marks_list:
                if row_idx >= total_row:
                    break
                try:
                    marks_table.rows[row_idx].cells[col_idx].text = str(mark)
                    total += int(mark)
                except:
                    pass
                row_idx += 1

            marks_table.rows[total_row].cells[col_idx].text = str(total)

        # -------------------------------------------------
        # Safe save
        # -------------------------------------------------
        try:
            doc.save(output_path)
        except PermissionError:
            ts = time.strftime("%Y%m%d_%H%M%S")
            alt = os.path.join(
                output_dir,
                f"{team_name_fs}_Review3_{ts}.docx"
            )
            doc.save(alt)
            output_path = alt

        return JsonResponse(
            {
                "status": "success",
                "message": "Review 2 marks saved successfully",
                "file_path": output_path
            }
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse(
            {"status": "error", "message": str(e)},
            status=500
        )
from django.views.decorators.csrf import csrf_exempt
from django.utils import timezone

@csrf_exempt
def request_reupload(request):
    """Mentor requests student to re-upload files"""
    mentor_name = request.session.get("mentor_name")
    username = request.session.get("username")
    
    if not mentor_name:
        return JsonResponse({"status": "fail", "message": "Not logged in"}, status=401)
    
    if request.method != "POST":
        return JsonResponse({"status": "fail", "message": "POST only"}, status=405)
    
    try:
        data = json.loads(request.body.decode('utf-8'))
        team_name = data.get("team_name")
        reason = data.get("reason", "Please re-upload your files.")
        
        allocation = AllocationResult.objects.filter(
            mentor_name=mentor_name, 
            team_name=team_name
        ).first()
        
        if not allocation:
            return JsonResponse({"status": "fail", "message": "Team not found"}, status=404)
        
        # Update Team model
        team = Team.objects.filter(project_title=team_name).first()
        if team:
            team.reupload_allowed = True
            team.reupload_reason = reason
            team.reupload_requested_at = timezone.now()
            team.reupload_requested_by = mentor_name
            team.save()
            
            # 🔥 Delete existing ProjectFile records so student can re-upload
            ProjectFile.objects.filter(team_name=team_name, review_type="zero").delete()
            
            return JsonResponse({
                "status": "success", 
                "message": "Re-upload requested. Student can now upload again."
            })
        else:
            return JsonResponse({"status": "fail", "message": "Team not found"}, status=404)
            
    except Exception as e:
        return JsonResponse({"status": "fail", "message": str(e)}, status=500)